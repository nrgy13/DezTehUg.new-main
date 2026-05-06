"""
Конвертирует 3 исходных DOCX из DTUdocs/ в docxtemplater-шаблоны.

Источники -> шаблоны:
- act-work-oliva-park.docx -> templates/act_work-dtudocs.docx
- ds4-arkhipo.docx -> templates/addendum-dtudocs.docx
- act-inspection-template.docx -> templates/act_inspection-dtudocs.docx

Стратегия замены:
1. Открываем DOCX через python-docx.
2. Идём по всем paragraphs (включая внутри таблиц) и применяем замены.
3. Для split-runs склеиваем текст параграфа, делаем замену, очищаем все runs
   кроме первого, в первый записываем результат (теряем inline форматирование
   внутри параграфа — для тегов docxtemplater это и нужно: каждый {tag}
   обязан лежать в одном <w:t>).
4. Для таблиц — то же по всем ячейкам.

Запуск:
    python tools/build-templates-from-dtudocs.py

Зависимости:
    pip install python-docx
"""

import os
import sys
import shutil

try:
    from docx import Document
except ImportError:
    print('ERROR: install python-docx -> pip install python-docx', file=sys.stderr)
    sys.exit(1)

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DOCX_DIR = os.path.join(ROOT, 'DTUdocs')
OUT_DIR = os.path.join(ROOT, 'tmp', 'templates-built')
os.makedirs(OUT_DIR, exist_ok=True)


SPACE_VARIANTS = [' ', ' ', ' ', ' ', ' ', ' ']


def normalize(text):
    """Заменяет non-breaking / narrow-no-break / thin spaces на обычные."""
    for s in SPACE_VARIANTS[1:]:
        text = text.replace(s, ' ')
    return text


def apply_replacements(text, replacements):
    """Применяет список замен с нормализацией пробелов в нужных местах."""
    norm = normalize(text)
    changed = False
    for old, new in replacements:
        # Сначала пробуем точное совпадение
        if old in norm:
            norm = norm.replace(old, new)
            changed = True
            continue
        # Потом нормализованный вариант
        old_norm = normalize(old)
        if old_norm in norm:
            norm = norm.replace(old_norm, new)
            changed = True
    return norm if changed else text


def replace_paragraph_text(paragraph, replacements):
    """Применяет замены к параграфу, склеивая runs."""
    full = paragraph.text
    new = apply_replacements(full, replacements)
    if new == full:
        return False
    runs = paragraph.runs
    if not runs:
        return False
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    runs[0].text = new
    return True


def replace_cell_text(cell, replacements):
    """Применяет замены к ВСЕЙ ячейке таблицы как единому блоку текста.

    Полезно когда значение разбито на несколько параграфов или runs внутри
    ячейки. Объединяет всё в первый параграф, теряя внутреннее форматирование.
    """
    full = '\n'.join(p.text for p in cell.paragraphs)
    new = apply_replacements(full, replacements)
    if new == full:
        return False
    # Удаляем все параграфы кроме первого
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    # Очищаем первый параграф
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(new)
    return True


def apply_to_document(doc, replacements):
    """Сначала параграфы документа (вне таблиц), потом ячейки таблиц целиком."""
    changed = 0
    for p in doc.paragraphs:
        if replace_paragraph_text(p, replacements):
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if replace_cell_text(cell, replacements):
                    changed += 1
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            if replace_cell_text(nc, replacements):
                                changed += 1
    return changed


# ─────────────────────────────────────────────────────────────────────
# 1. ACT WORK — на базе «Олива Парк»
# ─────────────────────────────────────────────────────────────────────
def build_act_work():
    src = os.path.join(
        DOCX_DIR,
        'Акт по проведению работ  ИП Белавина Санаторий Каникулы в Анапе Столовая Олива Парк.docx',
    )
    dst = os.path.join(OUT_DIR, 'act_work-dtudocs.docx')
    shutil.copy2(src, dst)
    doc = Document(dst)

    replacements = [
        # Шапка
        ('07.04.2026', '{document.date}'),
        ('№ _____', '№ {document.number}'),
        # Заказчик
        ('ООО «Аппетит»', '{client.shortName}'),
        ('2320155529', '{client.inn}'),
        ('Анна (Управляющая)', '{contact.fio}'),
        ('8-918-007-23-37', '{contact.phone}'),
        # Исполнитель — ОСТАВЛЯЕМ как есть (это ИП Белавина, фиксированный)
        # Кроме дезинфектора:
        ('Денисов Ю.Л.', '{master.fio}'),
        # Объект
        (
            'Санаторий «Каникулы в Анапе»/Столовая «Олива Парк»',
            '{object.name}',
        ),
        ('г. Анапа, Пионерский проспект,23', '{object.address}'),
        ('1 174,8 м2', '{object.area} м2'),
        # Услуга
        ('Дезинсекция (уничтожение тараканов)', '{services}'),
        # Договор-родитель
        ('№ДТЮ-28/01/26-16', '№{deal.contractNumber}'),
        ('«28 января» 2026г', '«{deal.contractDateLong}»'),
    ]
    changed = apply_to_document(doc, replacements)
    doc.save(dst)
    print(f'  act_work: {changed} замен -> {dst}')
    return dst


# ─────────────────────────────────────────────────────────────────────
# 2. ADDENDUM (ДС) — на базе ДС№4 (Архипо-Осиповка)
# ─────────────────────────────────────────────────────────────────────
def build_addendum():
    src = os.path.join(DOCX_DIR, 'ДС№4 ООО Аппетит (Архипо-Осиповка).docx')
    dst = os.path.join(OUT_DIR, 'addendum-dtudocs.docx')
    shutil.copy2(src, dst)
    doc = Document(dst)

    replacements = [
        # Заголовок
        ('Дополнительное соглашение №4', 'Дополнительное соглашение №{addendum.number}'),
        # Договор-родитель
        ('№ДТЮ-28/01/26-16', '№{deal.contractNumber}'),
        ('от 28.01.2026 г.', 'от {deal.contractDate} г.'),
        # Метаданные ДС
        ('Дата заключения: 23 апреля 2026 г', 'Дата заключения: {addendum.dateLong} г'),
        ('Место заключения: г. Новороссийск', 'Место заключения: {addendum.location}'),
        # Багфикс «Приложение №2» -> правильный номер
        ('Приложение №2', 'Приложение №{addendum.number}'),
        # Заказчик в преамбуле
        (
            'Общества с ограниченной ответственностью «Аппетит»',
            '{client.fullName}',
        ),
        ('действующего на основании Устава', 'действующего на основании {client.actingBasis}'),
        # Объект (одна строка-заглушка для демо; для нескольких объектов в ДС
        # потребуется вручную обернуть в {#objects}…{/objects})
        (
            'База отдыха «Рассвет, Краснодарский край, муниципальное образование Геленджик, село Архипо-Осиповка',
            '{object.name}, {object.address}',
        ),
        # Прайс — одну строку-данные превращаем в loop-теги.
        # В таблице 2 строки (дезинс. + дератиз.) — заменяю обе на ОДНУ
        # с тегами {#priceItems}…{/priceItems}.
        # Шаблонизатор docxtemplater дублирует TR для каждого элемента.
        # Стратегия: первую строку -> теги loop, вторую строку -> удалить (в build).
        ('Дезинсекция (уничтожение тараканов)', '{customName}'),
        ('Дератизация (уничтожение грызунов) пест-контроль', '__DELETE_ROW__'),
        ('530', '{areaM2}'),
        ('Сухая/Точечное орошение/Туман', '{method}'),
        ('По заявке', '{frequency}'),
        ('5 714,29', '{priceNoVat}'),
        ('6 000,00', '{priceWithVat}'),
        # Реквизиты заказчика (низ ДС)
        ('ООО «Аппетит»', '{client.shortName}'),
        (
            '354340, Краснодарский край, г. Сочи, ул. Гастелло, дом 28',
            '{client.legalAddress}',
        ),
        (
            '354000, Краснодарский край, г. Сочи, а/я 184',
            '{client.postalAddress}',
        ),
        ('2320155529', '{client.inn}'),
        ('1072320016801', '{client.ogrn}'),
        ('236701001', '{client.kpp}'),
        ('40702810030060009260', '{client.bankAccount}'),
        ('В ЮГО-ЗАПАДНОМ БАНКЕ ПАО СБЕРБАНК', 'В {client.bankName}'),
        ('046015602', '{client.bankBik}'),
        ('30101810600000000602', '{client.bankCorrAccount}'),
        ('8-988-236-05-07', '{client.phone}'),
        ('info@appetit.su', '{client.email}'),
        # Подпись заказчика
        ('Генеральный директор', '{client.directorRole}'),
        ('А.Е.Мороз', '{client.directorShort}'),
    ]
    changed = apply_to_document(doc, replacements)

    # Удаляем строки таблицы где встретился маркер __DELETE_ROW__
    deleted_rows = 0
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            row_text = ' '.join(p.text for cell in row.cells for p in cell.paragraphs)
            if '__DELETE_ROW__' in row_text:
                rows_to_remove.append(row)
        for r in rows_to_remove:
            r._element.getparent().remove(r._element)
            deleted_rows += 1

    doc.save(dst)
    print(f'  addendum: {changed} замен, {deleted_rows} удалённых строк -> {dst}')
    return dst


# ─────────────────────────────────────────────────────────────────────
# 3. ACT INSPECTION — шаблон уже почти пустой, добавим только number/date
# ─────────────────────────────────────────────────────────────────────
def build_act_inspection():
    src = os.path.join(DOCX_DIR, 'Акт обледования шаблон ИП Белавина.docx')
    dst = os.path.join(OUT_DIR, 'act_inspection-dtudocs.docx')
    shutil.copy2(src, dst)
    doc = Document(dst)

    replacements = [
        # Шапка с датой
        ('АКТ-ОБСЛЕДОВАНИЯ от ', 'АКТ-ОБСЛЕДОВАНИЯ № {document.number} от {document.date} '),
    ]
    # В этом шаблоне нет заполненных данных клиента/объекта — Регина пишет
    # руками после распечатки. Поэтому замен немного.
    changed = apply_to_document(doc, replacements)
    doc.save(dst)
    print(f'  act_inspection: {changed} замен -> {dst}')
    return dst


if __name__ == '__main__':
    print('Building docxtemplater templates from DTUdocs...')
    build_act_work()
    build_addendum()
    build_act_inspection()
    print(f'\nDone. Templates in: {OUT_DIR}')
