# -*- coding: utf-8 -*-
"""
Специализированный парсер договоров Регины (DTU_docs_clients).

Парсит структурированно — по конкретным таблицам и параграфам, а не regex'ами
по слитому тексту. Покрывает 3 типа файлов:

  * MAIN  — основной договор (~31 файл, 6 таблиц): Table 1 = реквизиты (исполнитель|заказчик),
            Table N = прайс с заголовками «Наименование услуги|Объект|Площадь|Стоимость|Периодичность»,
            параграф перед прайсом = список объектов обработки.
  * ADDENDUM — Доп.соглашение (ДС) (~9 файлов, 3-4 таблицы): меньше структура,
               только реквизиты + изменения. Привязывается к основному договору по ИНН.
  * UNKNOWN — нестандартные форматы. Парсим best-effort, помечаем для ручной проверки.

Выход:
  tmp/regina-contracts-v2.json — список объектов:
    {
      type: "main"|"addendum"|"unknown",
      filename, file_type, contract: {number, date, place},
      client: { все реквизиты заказчика },
      objects: [ {name, address} ],
      price_items: [ {service, object, area_m2, price_no_vat, frequency} ],
      raw_warnings: [...]   // что не удалось распарсить
    }

  tmp/regina-clients.json — сгруппировано по ИНН:
    { "1831202576": { client_data, contracts: [...], all_price_items: [...], objects: [...] } }

Запуск:
    python3 tools/parse-regina-contracts.py                      # все 47 файлов
    python3 tools/parse-regina-contracts.py path/to/file.docx    # один файл, печать
"""
import os
import sys
import re
import json
from glob import glob
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print('ERROR: pip install python-docx', file=sys.stderr)
    sys.exit(1)

# Морфологический анализатор для именительного падежа ФИО директоров.
# Договоры пишут в винительном/род.п.: «в лице директора Башкировой Татьяны Юрьевны»,
# но в БД нужно «Башкирова Татьяна Юрьевна».
try:
    import pymorphy3
    _MORPH = pymorphy3.MorphAnalyzer()
except ImportError:
    _MORPH = None


def to_nominative(word):
    """Возвращает слово в именительном падеже (через pymorphy3). Сохраняет регистр."""
    if not _MORPH or not word:
        return word
    parsed = _MORPH.parse(word)
    if not parsed:
        return word
    nomn = parsed[0].inflect({'nomn'})
    if not nomn:
        return word
    out = nomn.word
    # Восстановим капитализацию (первая буква большая если у исходного была)
    if word[0].isupper():
        out = out[0].upper() + out[1:]
    return out


def _detect_gender(words):
    """Пол по отчеству (последнее слово ФИО). Нужен чтобы корректно склонять фамилию:
    «Мирошева Алексея Владимировича» — мужчина (Владимир-ВИЧ), значит «Мирошева» это
    род.п. от «Мирошев», а НЕ им.п. женской фамилии. Без пола pymorphy ошибается."""
    if not words:
        return None
    patr = words[-1].lower()
    if 'вич' in patr:                       # Владимирович/а/у/ем
        return 'masc'
    if 'вн' in patr or 'чн' in patr:        # Юрьевна/ы, Ильинична
        return 'femn'
    return None


def _word_nomn(word, prefer_tags=(), gender=None):
    """Слово → им.п. Предпочитает разборы с нужными граммемами (Surn/Name/Patr) и полом."""
    if not _MORPH or not word:
        return word
    variants = _MORPH.parse(word)
    if not variants:
        return word
    cand = variants
    if prefer_tags:
        tagged = [v for v in variants if any(t in v.tag for t in prefer_tags)]
        if tagged:
            cand = tagged
    if gender:
        gendered = [v for v in cand if gender in v.tag]
        if gendered:
            cand = gendered
    # Слово уже в им.п. ед.ч. (несклоняемые фамилии «Шупик», «Азатян») — не трогаем
    for v in cand:
        if 'nomn' in v.tag and 'sing' in v.tag:
            return word
    best = cand[0]
    nomn = best.inflect({'nomn'})
    out = nomn.word if nomn else word
    if word and word[0].isupper():
        out = out[0].upper() + out[1:]
    return out


def fio_to_nominative(fio):
    """«Мирошева Алексея Владимировича» → «Мирошев Алексей Владимирович».
    Фамилию (1-е слово) склоняем с учётом пола по отчеству; имя/отчество — по своим граммемам.
    Также используется для роли («Генерального директора» → «Генеральный директор»)."""
    if not fio:
        return fio
    words = fio.split()
    gender = _detect_gender(words)
    out = []
    n = len(words)
    for i, w in enumerate(words):
        if i == 0 and n >= 2:
            out.append(_word_nomn(w, prefer_tags=('Surn',), gender=gender))
        elif i == n - 1 and n >= 3:
            out.append(_word_nomn(w, prefer_tags=('Patr',), gender=gender))
        elif n >= 3:
            out.append(_word_nomn(w, prefer_tags=('Name',), gender=gender))
        else:
            out.append(_word_nomn(w))
    return ' '.join(out)


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOURCE_DIR = os.path.join(ROOT, 'DTU_docs_clients', 'Договора на заключение — копия')

# Реквизиты исполнителя — Регина
EXECUTOR_INN = '231507022304'
EXECUTOR_OGRNIP = '321237500467390'

# ─── Утилиты ──────────────────────────────────────────────────────────


def nbsp(s):
    """Уберёт NBSP и nbsp-варианты пробелов."""
    if s is None:
        return ''
    s = s.replace('\xa0', ' ').replace(' ', ' ')  # narrow no-break space
    s = re.sub(r'[ \t]+', ' ', s)
    return s.strip()


def parse_money(s):
    """'10 900,48' → 10900.48 ; '4500,00' → 4500.00 ; '-' → None"""
    if not s:
        return None
    raw = nbsp(s)
    if raw in ('-', '—', '–', ''):
        return None
    cleaned = raw.replace(' ', '').replace(',', '.')
    cleaned = re.sub(r'[^\d.]', '', cleaned)
    try:
        return float(cleaned)
    except ValueError:
        return None


def parse_area(s):
    """
    '138,20' → 138.2
    '2 081'  → 2081
    'до 30'  → 30
    '-'      → None
    'Весь объект' → None (но возвращаем строкой через area_label)
    """
    if not s:
        return None, None
    raw = nbsp(s)
    if raw in ('-', '—', '–', ''):
        return None, None
    m = re.search(r'(\d+(?:[,.]\d+)?)', raw.replace(' ', ''))
    if m:
        try:
            return float(m.group(1).replace(',', '.')), raw
        except ValueError:
            return None, raw
    # Не нашли число — это «Весь объект», «вся территория» и т.п.
    return None, raw


MONTHS_RU = {
    'январ': 1, 'феврал': 2, 'март': 3, 'апрел': 4,
    'мая': 5, 'мае': 5, 'май': 5,
    'июн': 6, 'июл': 7, 'август': 8,
    'сентябр': 9, 'октябр': 10, 'ноябр': 11, 'декабр': 12,
}


def ru_date_to_iso(day, month_word, year):
    month_word = month_word.lower()
    for key, num in MONTHS_RU.items():
        if month_word.startswith(key):
            return f"{int(year):04d}-{num:02d}-{int(day):02d}"
    return None


# ─── Идентификация типа файла ─────────────────────────────────────────


def classify_file(fname):
    """main / addendum / unknown."""
    base = os.path.basename(fname).lower()
    if base.startswith('дс') or base.startswith('д.с') or 'дс №' in base or 'дс №' in base or 'дс№' in base:
        return 'addendum'
    if base.startswith('договор') or base.startswith('контракт') or 'договор от' in base:
        return 'main'
    return 'unknown'


def parse_filename_meta(fname):
    """
    Возвращает (date_iso_or_none, hint_org_type, hint_short_name, hint_object).
    Берёт всё что можно из имени файла как fallback к парсингу контента.
    """
    base = os.path.basename(fname).rsplit('.', 1)[0]
    base = re.sub(r'\s+', ' ', base).strip()
    # Дата dd.mm.yyyy
    m = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', base)
    date_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}" if m else None
    # Чистим: убираем "ДОГОВОР"/"Договор" + "ДС"/"ДС№N" + "от <date>"
    cleaned = re.sub(r'^(Договор|ДОГОВОР|Контракт|ДС\s*№?\s*\d*|ДС|Д\.\s*С\.?)\s*', '', base, flags=re.IGNORECASE)
    cleaned = re.sub(r'\s*от\s+\d{2}\.\d{2}\.\d{4}\s*г?\.?\s*', ' ', cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r'\s+', ' ', cleaned)
    # Тип юр.лица
    m = re.match(r'^(ООО|ИП|АО|ПАО|ЗАО|ОАО)\s+(.+?)$', cleaned)
    if not m:
        return date_iso, None, None, None
    org_type, rest = m.group(1), m.group(2).strip()
    obj_hint = None
    paren = re.search(r'^(.+?)\s*\(([^)]+)\)\s*(.*)$', rest)
    if paren:
        short_name = paren.group(1).strip()
        obj_hint = paren.group(2).strip()
    else:
        # «г. <место>» в конце — это объект
        m2 = re.split(r'\s+(?:г\.|пос?\.|с\.|ст\.|х\.|мкр\.)\s+', rest, maxsplit=1)
        if len(m2) == 2:
            short_name = m2[0].strip()
            obj_hint = m2[1].strip()
        else:
            short_name = rest.strip()
    return date_iso, org_type, short_name, obj_hint


# ─── Идентификация таблиц по содержимому ──────────────────────────────


def is_requisites_table(t):
    """Таблица реквизитов: 1 строка × 2 колонки, содержит ИНН + ОГРН/ОГРНИП."""
    if len(t.rows) != 1 or len(t.columns) != 2:
        return False
    text = ' '.join(c.text for c in t.rows[0].cells)
    return 'ИНН' in text and ('ОГРН' in text or 'ОГРНИП' in text)


def is_price_table(t):
    """Таблица прайса: заголовок содержит «услуг» + «стоимост»|«цена»."""
    if len(t.rows) < 2:
        return False
    header = ' '.join(nbsp(c.text).lower() for c in t.rows[0].cells)
    has_service = 'услуг' in header or 'наименован' in header
    has_price = 'стоимост' in header or 'цена' in header or 'тариф' in header
    return has_service and has_price


def is_application_form_table(t):
    """Заявка (Приложение №1): заголовок «Дата | Время | Место | Объем»."""
    if len(t.rows) < 1:
        return False
    header = ' '.join(nbsp(c.text).lower() for c in t.rows[0].cells)
    return 'место обработки' in header and ('дата' in header or 'время' in header)


# ─── Парсинг ячейки реквизитов ────────────────────────────────────────


def parse_requisites_cell(text):
    """
    Парсит одну ячейку с реквизитами стороны (Исполнитель ИЛИ Заказчик).
    Текст идёт построчно — заголовок «ИСПОЛНИТЕЛЬ:» / «ЗАКАЗЧИК:», название,
    адреса, ИНН, ОГРН, расчётный счёт, банк, БИК, корр.счёт, тел, email.
    """
    out = {
        'role_header': None,
        'name_raw': None,
        'org_type': None,
        'short_name': None,
        'full_name': None,
        'legal_address': None,
        'actual_address': None,
        'inn': None,
        'kpp': None,
        'ogrn': None,
        'ogrnip': None,
        'bank_account': None,
        'bank_name': None,
        'bank_bik': None,
        'bank_corr_account': None,
        'phone': None,
        'email': None,
    }
    lines = [nbsp(ln) for ln in text.split('\n') if nbsp(ln)]
    if not lines:
        return out

    # 1. Шапка стороны
    if lines[0].upper().startswith(('ИСПОЛНИТЕЛЬ', 'ЗАКАЗЧИК')):
        out['role_header'] = lines[0].rstrip(':')
        body_lines = lines[1:]
    else:
        body_lines = lines

    # 2. Название (первая непустая строка после шапки, до «Адрес»/«ИНН»)
    name_lines = []
    body_idx = 0
    while body_idx < len(body_lines):
        ln = body_lines[body_idx]
        if re.match(r'^(Юридический|Фактический|Почтовый|Адрес|ИНН|КПП|ОГРН|Расчетн|БИК|Тел|E[\s\-]?mail)', ln, re.IGNORECASE):
            break
        name_lines.append(ln)
        body_idx += 1
    if name_lines:
        out['name_raw'] = ' '.join(name_lines).strip()
        # Тип организации
        m = re.match(r'^(ООО|АО|ПАО|ЗАО|ОАО)\s*[«"\']?([^»"\']+)[»"\']?', out['name_raw'])
        if m:
            out['org_type'] = m.group(1)
            short = m.group(2).strip().strip('«»"\'')
            out['short_name'] = f'{m.group(1)} «{short}»'
        elif re.match(r'^(ИП|Индивидуальн)', out['name_raw'], re.IGNORECASE):
            # ИП Фамилия И.О.
            mm = re.match(r'(?:ИП|Индивидуальн\w+\s+предпринимател\w+)\s+(.+?)$', out['name_raw'], re.IGNORECASE)
            if mm:
                out['org_type'] = 'ИП'
                out['short_name'] = f'ИП {mm.group(1).strip()}'
        elif re.match(r'^Общество', out['name_raw'], re.IGNORECASE):
            mm = re.search(r'Общество с ограниченной ответственностью\s*[«"\']?([^»"\']+)[»"\']?', out['name_raw'], re.IGNORECASE)
            if mm:
                out['org_type'] = 'ООО'
                short = mm.group(1).strip().strip('«»"\'')
                out['short_name'] = f'ООО «{short}»'
                out['full_name'] = out['name_raw']

    # 3. Поля построчно
    rest = body_lines[body_idx:]
    bank_name_pending = False
    for ln in rest:
        # Адреса
        m = re.match(r'Юридический\s+адрес[:\s]+(.+?)\s*$', ln, re.IGNORECASE)
        if m:
            out['legal_address'] = m.group(1).strip()
            # «Юридический ... Фактический ...» может быть в одной строке
            m2 = re.search(r'Фактический\s+адрес[:\s]+(.+?)$', out['legal_address'], re.IGNORECASE)
            if m2:
                out['actual_address'] = m2.group(1).strip()
                out['legal_address'] = re.sub(r'\s*Фактический\s+адрес.*$', '', out['legal_address'], flags=re.IGNORECASE).strip()
            continue
        m = re.match(r'Фактический\s+адрес[:\s]+(.+?)\s*$', ln, re.IGNORECASE)
        if m:
            out['actual_address'] = m.group(1).strip()
            continue
        m = re.match(r'Почтовый\s+адрес[:\s]+(.+?)\s*$', ln, re.IGNORECASE)
        if m:
            out.setdefault('postal_address', m.group(1).strip())
            continue
        # ИНН — 10 (юр) или 12 (ИП)
        m = re.search(r'ИНН[:\s]*(\d{10,12})', ln)
        if m and not out['inn']:
            out['inn'] = m.group(1)
        m = re.search(r'КПП[:\s]*(\d{9})', ln)
        if m and not out['kpp']:
            out['kpp'] = m.group(1)
        m = re.search(r'ОГРНИП[:\s]*(\d{15})', ln)
        if m and not out['ogrnip']:
            out['ogrnip'] = m.group(1)
        else:
            m = re.search(r'ОГРН[:\s]*(\d{13})', ln)
            if m and not out['ogrn']:
                out['ogrn'] = m.group(1)
        # Расчётный счёт
        m = re.search(r'Расчетн\w+\s+счет[:\s]*(\d{20})', ln, re.IGNORECASE)
        if m and not out['bank_account']:
            out['bank_account'] = m.group(1)
            bank_name_pending = True   # следующая непустая строка = название банка
            continue
        # Название банка идёт сразу после расчётного счёта
        if bank_name_pending and not re.match(r'^(БИК|Корр|Тел|E[\s\-]?mail|КПП|ОГРН|ИНН)', ln, re.IGNORECASE):
            out['bank_name'] = ln.strip()
            bank_name_pending = False
            continue
        bank_name_pending = False
        m = re.search(r'БИК[:\s]*(\d{9})', ln)
        if m and not out['bank_bik']:
            out['bank_bik'] = m.group(1)
        m = re.search(r'Корр[.\s]*счет[:\s]*(\d{20})', ln, re.IGNORECASE)
        if m and not out['bank_corr_account']:
            out['bank_corr_account'] = m.group(1)
        m = re.search(r'Тел[.\s:]+([+\d\-()\s.]{10,30})', ln, re.IGNORECASE)
        if m and not out['phone']:
            out['phone'] = nbsp(m.group(1))
        m = re.search(r'[Ee][\s\-]?mail[:\s]*([A-Za-z0-9_.+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]+)', ln)
        if m and not out['email']:
            out['email'] = m.group(1)

    return out


# ─── Парсинг прайса ───────────────────────────────────────────────────


SERVICE_CODE_MAP = [
    (re.compile(r'дезинсекц', re.IGNORECASE), 'disinsection'),
    (re.compile(r'дератизац|пест[- ]контрол|грызун', re.IGNORECASE), 'deratization'),
    (re.compile(r'дезинфекц', re.IGNORECASE), 'disinfection'),
    (re.compile(r'фумигац', re.IGNORECASE), 'fumigation'),
    (re.compile(r'дезодорац|озонировани', re.IGNORECASE), 'deodorization'),
    (re.compile(r'десерпентац|зме', re.IGNORECASE), 'deserpentation'),
    (re.compile(r'гербицид|расти?тельност', re.IGNORECASE), 'herbicide-treatment'),
    (re.compile(r'анализ\s+воды', re.IGNORECASE), 'water-analysis'),
    (re.compile(r'москит|комар|клещ|муравь', re.IGNORECASE), 'disinsection'),
]


def service_to_code(name):
    """Маппит наименование услуги на код из каталога services (см. lib/db/seed.ts)."""
    if not name:
        return None
    for pat, code in SERVICE_CODE_MAP:
        if pat.search(name):
            return code
    return None


def parse_price_table(t):
    """
    Парсит таблицу прайса.
    Стандартный заголовок (5 колонок): Наименование услуги | Объект | Площадь м² | Стоимость | Периодичность
    Бывает другой порядок — определяем по заголовку.
    Также фильтруем «разделительные» строки типа «Прочие услуги|Прочие услуги|...».
    """
    if len(t.rows) < 2:
        return []
    header = [nbsp(c.text).lower() for c in t.rows[0].cells]

    def find_col(*keys):
        for i, h in enumerate(header):
            if any(k in h for k in keys):
                return i
        return None

    i_service = find_col('наименован', 'услуг')
    i_object = find_col('объект', 'место', 'помещен')
    i_area = find_col('площад', 'м2', 'м²')
    i_price = find_col('стоимост', 'цена', 'тариф')
    i_freq = find_col('периодичн', 'кратност', 'графи')

    items = []
    for row in t.rows[1:]:
        cells = [nbsp(c.text) for c in row.cells]
        if not any(cells):
            continue
        # Разделители ("Прочие услуги|Прочие услуги|...")
        if len(set(cells)) == 1:
            continue
        # Пустые «комментарии к заявке»
        if cells[0].startswith('Комментар'):
            continue

        service_name = cells[i_service] if i_service is not None and i_service < len(cells) else None
        object_name = cells[i_object] if i_object is not None and i_object < len(cells) else None
        area_raw = cells[i_area] if i_area is not None and i_area < len(cells) else None
        price_raw = cells[i_price] if i_price is not None and i_price < len(cells) else None
        freq = cells[i_freq] if i_freq is not None and i_freq < len(cells) else None

        if not service_name or not price_raw or not parse_money(price_raw):
            continue  # пустая или некорректная строка

        area_num, area_label = parse_area(area_raw or '')
        items.append({
            'service_name': service_name,
            'service_code': service_to_code(service_name),
            'object_name': object_name,
            'area_m2': area_num,
            'area_label': area_label,
            'price_no_vat': parse_money(price_raw),
            'frequency': freq,
        })
    return items


# ─── Парсинг объектов (адреса) из параграфа над прайсом ──────────────


ADDRESS_HINTS = ('г.', 'г ', 'ул.', 'пр.', 'просп.', 'пер.', 'д.', 'дом', 'п.', 'пос.', 'с.', 'ст.', 'мкр.', 'шоссе', 'наб.', 'пр-т', 'пр-кт')


def find_objects_paragraph(doc):
    """
    Список объектов лежит в Приложении №2 — между «СТОИМОСТЬ УСЛУГ» (заголовок
    приложения, все буквы заглавные) и таблицей прайса.

    КРИТИЧНО: НЕ путать с разделом 4 договора «Стоимость услуг и порядок расчётов» —
    он содержит юр.текст и идёт раньше. Ищем именно «СТОИМОСТЬ УСЛУГ» (uppercase)
    и/или «Приложение № 2» как якорь.
    """
    paragraphs = [nbsp(p.text) for p in doc.paragraphs]

    # Ищем якорь Приложения №2 (предпочтительнее) — он гарантированно в шапке прайса
    anchor_idx = None
    for i, p in enumerate(paragraphs):
        stripped = p.strip()
        # «Приложение № 2 к Договору №ДТЮ-...»
        if re.match(r'Приложение\s*№?\s*2\b', stripped, re.IGNORECASE):
            anchor_idx = i
            break

    # Fallback: только если параграф ЦЕЛИКОМ = "СТОИМОСТЬ УСЛУГ" (uppercase, без пунктов)
    if anchor_idx is None:
        for i, p in enumerate(paragraphs):
            stripped = p.strip().rstrip('.:')
            if stripped == 'СТОИМОСТЬ УСЛУГ':
                anchor_idx = i
                break

    if anchor_idx is None:
        return []

    # После якоря — собираем до подписей / следующего приложения / конца
    # И берём только параграфы которые ЛОГИЧНО являются адресами (содержат hints).
    objects = []
    seen = set()
    MAX_SCAN = 30  # объектов в одном договоре редко больше 5-6, 30 параграфов после якоря с запасом
    skip_titles = ('СТОИМОСТЬ УСЛУГ', 'от ', 'м.п.', 'Подпись', '____')
    for p in paragraphs[anchor_idx + 1:anchor_idx + 1 + MAX_SCAN]:
        if not p:
            continue
        if p.startswith(('*', 'ИП Белавина', 'Директор', 'Подпись', 'м.п.', '_____')):
            break  # пошли подписи
        if re.match(r'Приложение\s*№?\s*\d+', p, re.IGNORECASE):
            break  # следующее приложение
        # Заголовок «СТОИМОСТЬ УСЛУГ» / «от 23 января ...» — пропускаем
        if any(p.startswith(t) for t in skip_titles):
            continue
        if p.strip() == 'СТОИМОСТЬ УСЛУГ' or p.strip().startswith('от '):
            continue
        # Объекты разделяются «/» или \n (несколько в одном параграфе)
        for part in re.split(r'[/\n]', p):
            part = part.strip().rstrip(',').strip()
            if not part or len(part) < 5:
                continue
            # Дедупликация
            if part in seen:
                continue
            # Должен содержать адресный hint (иначе это просто текст)
            lower = part.lower()
            if not any(h in lower for h in ADDRESS_HINTS):
                continue
            seen.add(part)
            # name = первая фраза до запятой; address = остаток
            m = re.match(r'^([^,]+?)(?:,\s*(.+))?$', part)
            if m:
                name = m.group(1).strip()
                addr = (m.group(2) or '').strip()
                objects.append({'name': name, 'address': addr or None})
    return objects


# ─── Парсинг контракта (номер, дата, место, директор) ────────────────


def parse_contract_meta(doc):
    """
    Из первых параграфов:
      P0: «ДОГОВОР НА ОКАЗАНИЕ УСЛУГ №ДТЮ-15/01/26-14»
      P2: «Дата заключения: 23 января 2026 г»
      P3: «Место заключения: г. Новороссийск»
      P7: «и Общество ... в лице директора Башкировой Татьяны Юрьевны, действующей на основании Устава»
    """
    paragraphs = [nbsp(p.text) for p in doc.paragraphs if nbsp(p.text)]
    text_top = '\n'.join(paragraphs[:25])

    number = None
    date_iso = None
    place = None
    director_name = None
    director_role = None
    acting_basis = None

    m = re.search(r'№\s*((?:ДТЮ|ДГ|ДЮ|ДТ)[\-/\d/]+)', text_top)
    if m:
        number = m.group(1).strip()

    m = re.search(r'Дата заключения:\s*(\d{1,2})\s+([а-яё]+)\s+(\d{4})', text_top, re.IGNORECASE)
    if m:
        date_iso = ru_date_to_iso(m.group(1), m.group(2), m.group(3))

    m = re.search(r'Место заключения:\s*([^\n]+)', text_top, re.IGNORECASE)
    if m:
        place = m.group(1).strip().rstrip('.')

    # ── Директор ЗАКАЗЧИКА ──────────────────────────────────────────────
    # КРИТИЧНО: первый блок преамбулы = Исполнитель (ИП Белавина О.В., Регина):
    #   «...в лице Индивидуального предпринимателя Белавиной Ольги Владимировны
    #    (далее – Исполнитель), действующий на основании лицензии...».
    # Заказчик идёт ПОСЛЕ «..., с одной стороны, и ...». Отрезаем блок Исполнителя,
    # иначе ловим Белавину вместо реального заказчика.
    ogrnip = None
    zone = text_top
    m_split = re.search(r'с одной стороны', text_top)
    if m_split:
        zone = text_top[m_split.end():]
    else:
        m_exec = re.search(r'Исполнител', text_top)
        if m_exec:
            zone = text_top[m_exec.end():]

    # Вариант ООО: «в лице [Генерального ]директора ФИО(род.п.), действующ... на основании Устава».
    # Роль может начинаться с заглавной («Генерального директора») — раньше regex её терял.
    m = re.search(
        r'в лице\s+((?:[А-Яа-яё][а-яё]+\s+)*директор\w*)\s+'
        r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?)'
        r'\s*,?\s*действующ\w+\s+на\s+основании\s+([А-Яа-яё\s]+?)(?:,|\.|$)',
        zone,
    )
    if m:
        # Роль и ФИО заказчика-ООО в род./вин. падеже → приводим к именительному.
        role = fio_to_nominative(m.group(1).strip())
        director_role = (role[0].upper() + role[1:]) if role else role
        director_name = fio_to_nominative(m.group(2).strip())
        acting_basis = m.group(3).strip().rstrip('.')
    else:
        # Вариант ИП: «и Индивидуальный предприниматель ФИО(им.п.), действующ...
        #   на основании ОГРНИП: NNN». ФИО уже в ИМЕНИТЕЛЬНОМ — НЕ склоняем!
        m = re.search(
            r'Индивидуальн\w+\s+предпринимател\w+\s+'
            r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)'
            r'\s*,?\s*действующ\w+\s+на\s+основании\s+([^,\.]+)',
            zone,
        )
        if m:
            director_role = 'Индивидуальный предприниматель'
            director_name = m.group(1).strip()
            acting_basis = m.group(2).strip().rstrip('.')
            mo = re.search(r'ОГРНИП[:\s]*(\d{13,15})', m.group(2))
            if mo:
                ogrnip = mo.group(1)

    return {
        'number': number,
        'date': date_iso,
        'place': place,
        'director_name': director_name,
        'director_role': director_role,
        'acting_basis': acting_basis,
        'ogrnip': ogrnip,
    }


# ─── Главный парсинг ──────────────────────────────────────────────────


def parse_one(docx_path):
    warnings = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        return {'filename': os.path.basename(docx_path), 'error': str(e)}

    file_type = classify_file(docx_path)
    fname_date, fname_org_type, fname_short_name, fname_object = parse_filename_meta(docx_path)

    # 1. Реквизиты — таблица 1×2 с ИНН
    req_table = None
    for t in doc.tables:
        if is_requisites_table(t):
            req_table = t
            break
    client = {}
    executor = {}
    if req_table:
        cell_data = [parse_requisites_cell(c.text) for c in req_table.rows[0].cells]
        # Заказчик = та сторона у которой ИНН ≠ EXECUTOR_INN и ≠ EXECUTOR_OGRNIP
        for d in cell_data:
            if (d.get('inn') == EXECUTOR_INN) or (d.get('ogrnip') == EXECUTOR_OGRNIP):
                executor = d
            else:
                client = d
        # Если обе стороны не наши — берём вторую (по позиции в таблице)
        if not client and len(cell_data) == 2:
            client = cell_data[1]
    else:
        warnings.append('no_requisites_table')

    # Если из реквизитов не вышло имя — добиваем из имени файла
    if not client.get('short_name') and fname_short_name and fname_org_type:
        if fname_org_type == 'ООО':
            client['short_name'] = f'ООО «{fname_short_name}»'
            client['full_name'] = client.get('full_name') or f'Общество с ограниченной ответственностью «{fname_short_name}»'
        elif fname_org_type == 'ИП':
            client['short_name'] = f'ИП {fname_short_name}'
            client['full_name'] = client.get('full_name') or f'Индивидуальный предприниматель {fname_short_name}'
        else:
            client['short_name'] = f'{fname_org_type} «{fname_short_name}»'
        client['org_type'] = client.get('org_type') or fname_org_type

    # 2. Прайс — все таблицы которые pass is_price_table
    price_items = []
    for t in doc.tables:
        if is_price_table(t):
            price_items.extend(parse_price_table(t))

    # 3. Объекты — параграф над прайсом
    objects = find_objects_paragraph(doc)

    # 4. Контракт-мета
    contract_meta = parse_contract_meta(doc)
    if not contract_meta['date'] and fname_date:
        contract_meta['date'] = fname_date

    # ОГРНИП из преамбулы (заказчик-ИП) — если в реквизитах-таблице не нашёлся
    if not client.get('ogrnip') and contract_meta.get('ogrnip'):
        client['ogrnip'] = contract_meta['ogrnip']

    # 5. Подсказка из имени файла → доп. объект
    if fname_object and not objects:
        objects.append({'name': fname_object, 'address': None, 'source': 'filename'})

    # Warnings
    if not client.get('inn'):
        warnings.append('no_client_inn')
    if not price_items and file_type == 'main':
        warnings.append('no_price_items')

    return {
        'type': file_type,
        'filename': os.path.basename(docx_path),
        'contract': contract_meta,
        'client': client,
        'objects': objects,
        'price_items': price_items,
        'warnings': warnings,
    }


def main():
    args = sys.argv[1:]
    if args:
        result = parse_one(args[0])
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    files = sorted(glob(os.path.join(SOURCE_DIR, '*.docx')))
    # Исключаем отменённые договоры («ОТМЕНА ...») — клиент недействующий, не импортируем
    skipped = [f for f in files if os.path.basename(f).upper().startswith('ОТМЕНА')]
    files = [f for f in files if not os.path.basename(f).upper().startswith('ОТМЕНА')]
    print(f"Found {len(files)} .docx files (skipped {len(skipped)} ОТМЕНА)", file=sys.stderr)
    results = []
    for f in files:
        r = parse_one(f)
        results.append(r)

    # === Сохраняем сырые результаты ===
    os.makedirs(os.path.join(ROOT, 'tmp'), exist_ok=True)
    with open(os.path.join(ROOT, 'tmp', 'regina-contracts-v2.json'), 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"→ tmp/regina-contracts-v2.json ({len(results)} files)", file=sys.stderr)

    # === Группировка по ИНН ===
    by_inn = defaultdict(lambda: {'client': None, 'contracts': [], 'addendums': [], 'all_objects': [], 'all_price_items': []})
    no_inn_results = []
    for r in results:
        inn = r.get('client', {}).get('inn')
        if not inn:
            no_inn_results.append(r)
            continue
        grp = by_inn[inn]
        # client — берём из последнего (обычно самый полный)
        if not grp['client'] or len(str(r['client'])) > len(str(grp['client'])):
            grp['client'] = r['client']
        if r['type'] == 'addendum':
            grp['addendums'].append({'filename': r['filename'], 'contract': r['contract'], 'price_items': r['price_items'], 'objects': r['objects']})
        else:
            grp['contracts'].append({'filename': r['filename'], 'contract': r['contract']})
        grp['all_objects'].extend(r['objects'])
        grp['all_price_items'].extend(r['price_items'])

    # ── Фикс 3: привязка ДС-файлов без ИНН к клиентам по имени файла ──
    # Сопоставим имя файла ДС с уже найденными клиентами через короткое имя.
    def fname_keys(fname):
        """Достаёт ключевые слова из имени файла ДС для матчинга с клиентом."""
        base = fname.rsplit('.', 1)[0].lower()
        # Уберём префиксы ДС, ДС№N, "ИП", "ООО" и т.п.
        cleaned = re.sub(r'дс\s*№?\s*\d*\s*', '', base, flags=re.IGNORECASE)
        cleaned = re.sub(r'договор\s*', '', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'(ооо|ип|ао|пао)\s+', '', cleaned, flags=re.IGNORECASE)
        return set(re.findall(r'[а-яёa-z]{4,}', cleaned))  # слова длиной 4+

    orphan_addendums = []
    for r in no_inn_results:
        if r['type'] != 'addendum':
            orphan_addendums.append(r)
            continue
        keys = fname_keys(r['filename'])
        # Пытаемся найти клиента по совпадению ключевых слов в short_name
        best_inn = None
        best_score = 0
        for inn, grp in by_inn.items():
            client_name = (grp['client'].get('short_name') or '').lower()
            client_keys = set(re.findall(r'[а-яёa-z]{4,}', client_name))
            score = len(keys & client_keys)
            if score > best_score:
                best_score = score
                best_inn = inn
        if best_inn and best_score >= 1:
            by_inn[best_inn]['addendums'].append({
                'filename': r['filename'],
                'contract': r['contract'],
                'price_items': r['price_items'],
                'objects': r['objects'],
                'matched_by': 'filename_keywords',
            })
            by_inn[best_inn]['all_price_items'].extend(r['price_items'])
        else:
            orphan_addendums.append(r)

    # ── Фикс 2: fallback objects = legal_address если у клиента 0 объектов ──
    # Также генерируем full_name из short_name если не извлечён.
    for inn, grp in by_inn.items():
        client = grp['client']
        # Полное имя
        if not client.get('full_name') and client.get('short_name'):
            sn = client['short_name']
            if sn.startswith('ООО '):
                inner = sn.replace('ООО ', '').strip('«»"\' ')
                client['full_name'] = f'Общество с ограниченной ответственностью «{inner}»'
            elif sn.startswith('ИП '):
                fio = sn.replace('ИП ', '').strip()
                client['full_name'] = f'Индивидуальный предприниматель {fio}'
            else:
                client['full_name'] = sn
        # Объекты-fallback
        if not grp['all_objects'] and client.get('legal_address'):
            grp['all_objects'].append({
                'name': client.get('short_name') or 'Основной объект',
                'address': client['legal_address'],
                'source': 'fallback_legal_address',
            })
        # Дедупликация объектов
        seen = set()
        deduped = []
        for o in grp['all_objects']:
            key = ((o.get('name') or '').lower(), (o.get('address') or '').lower())
            if key in seen:
                continue
            seen.add(key)
            deduped.append(o)
        grp['all_objects'] = deduped

    grouped = {inn: data for inn, data in by_inn.items()}
    with open(os.path.join(ROOT, 'tmp', 'regina-clients.json'), 'w', encoding='utf-8') as f:
        json.dump(grouped, f, ensure_ascii=False, indent=2)
    print(f"→ tmp/regina-clients.json ({len(grouped)} unique clients by ИНН)", file=sys.stderr)
    if orphan_addendums:
        print(f"\nWARNING: {len(orphan_addendums)} files без ИНН и без матча по имени:", file=sys.stderr)
        for r in orphan_addendums:
            print(f"  • {r['filename']}", file=sys.stderr)

    # === Summary в txt ===
    with open(os.path.join(ROOT, 'tmp', 'regina-summary.txt'), 'w', encoding='utf-8') as f:
        f.write(f"=== ИТОГ: {len(results)} файлов → {len(grouped)} уникальных клиентов по ИНН ===\n")
        f.write(f"ДС-файлы без своего ИНН: {len(no_inn_results)} (из них orphan: {len(orphan_addendums)})\n\n")
        f.write(f"{'ИНН':>14} | объ | прайс | имя\n")
        f.write('-' * 100 + '\n')
        for inn, g in sorted(grouped.items(), key=lambda x: -len(x[1]['all_price_items'])):
            client = g['client'] or {}
            name = client.get('short_name', '???')
            n_obj = len(g['all_objects'])
            n_price = len(g['all_price_items'])
            n_addn = len(g['addendums'])
            addn = f' +{n_addn}ДС' if n_addn else ''
            f.write(f"  {inn:>14} | {n_obj:>3} | {n_price:>4}  | {name}{addn}\n")
    print(f"→ tmp/regina-summary.txt", file=sys.stderr)


if __name__ == '__main__':
    main()
