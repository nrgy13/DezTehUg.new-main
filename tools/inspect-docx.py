#!/usr/bin/env python3
"""Дамп структуры .docx для подготовки карт замен (шаблоны «точь-в-точь»).

Показывает: индекс/текст/число run'ов/флаг drawing для каждого параграфа,
содержимое всех таблиц (по ячейкам), суммарное число картинок (<a:blip>).
Run-дробление видно по [Nr] — если текст размазан по run'ам, set_merge их склеит.

ЗАПУСК: python tools/inspect-docx.py "<путь к .docx>"
"""
import sys
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass
from docx import Document
from docx.oxml.ns import qn


def has_drawing(run):
    return run._r.find(qn('w:drawing')) is not None or run._r.find(qn('w:pict')) is not None


def runs_repr(p):
    parts = []
    for r in p.runs:
        if has_drawing(r):
            parts.append("<IMG>")
        else:
            parts.append(repr(r.text))
    return " · ".join(parts)


def main(path):
    doc = Document(path)
    body = doc.element.body
    # blip count по всему документу
    xml = doc.element.xml
    blips = xml.count('<a:blip')
    print(f"=== {path}")
    print(f"=== Картинок (<a:blip>): {blips} | параграфов: {len(doc.paragraphs)} | таблиц: {len(doc.tables)}\n")

    print("──────── ПАРАГРАФЫ (тела документа) ────────")
    for i, p in enumerate(doc.paragraphs):
        txt = "".join(r.text for r in p.runs if not has_drawing(r))
        img = " [IMG]" if any(has_drawing(r) for r in p.runs) else ""
        nr = len(p.runs)
        if txt.strip() or img:
            print(f"[P{i:>3}] nr={nr}{img}  {txt!r}")
            if nr > 1 and txt.strip():
                print(f"        runs: {runs_repr(p)}")

    for ti, t in enumerate(doc.tables):
        print(f"\n──────── ТАБЛИЦА #{ti}  ({len(t.rows)}x{len(t.columns)}) ────────")
        for ri, row in enumerate(t.rows):
            cells = []
            for c in row.cells:
                ctxt = "\n".join("".join(r.text for r in pp.runs if not has_drawing(r)) for pp in c.paragraphs)
                cimg = "[IMG]" if any(has_drawing(r) for pp in c.paragraphs for r in pp.runs) else ""
                cells.append((cimg + ctxt).replace("\n", "⏎"))
            print(f"  R{ri}: " + " | ".join(cells))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("usage: python tools/inspect-docx.py <path.docx>"); sys.exit(1)
    main(sys.argv[1])
