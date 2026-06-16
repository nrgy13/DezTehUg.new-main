"""
Генератор всех DOCX-шаблонов CRM ДезТехЮг.

Запуск:
    python3 tools/build-templates.py            # все шаблоны
    python3 tools/build-templates.py contract   # только договор
    python3 tools/build-templates.py addendum   # только ДС
    python3 tools/build-templates.py inspection # только акт обследования
    python3 tools/build-templates.py work       # только акт работ
    python3 tools/build-templates.py offer      # только КП
    python3 tools/build-templates.py invoice    # только счёт

Placeholder-синтаксис: одинарные {tag} (стандарт docxtemplater).
Dot-notation поддерживается через парсер lib/render-docx.ts.

Маппинг полей описан в templates/README.md.
"""

from pathlib import Path
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path('templates')
OUT_DIR.mkdir(exist_ok=True)


# ─── Общие helpers ────────────────────────────────────────────


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    return p


def para(doc, text='', bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p


def fill_cell(cell, lines):
    """Заполняет ячейку: список (text, bold) → параграфы внутри."""
    cell.text = ''
    for i, (text, bold) in enumerate(lines):
        p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)


def requisites_table(doc):
    """Стандартная двухколоночная таблица реквизитов (Исполнитель | Заказчик)."""
    t = doc.add_table(rows=1, cols=2)
    fill_cell(t.rows[0].cells[0], [
        ('ИСПОЛНИТЕЛЬ:', True),
        ('{provider.shortName}', True),
        ('', False),
        ('Юридический адрес: {provider.legalAddress}', False),
        ('ИНН: {provider.inn}', False),
        ('ОГРНИП: {provider.ogrnip}', False),
        ('Расчётный счёт: {provider.bankAccount}', False),
        ('{provider.bankName}', False),
        ('БИК: {provider.bankBik}', False),
        ('Корр. счёт: {provider.bankCorrAccount}', False),
        ('Тел.: {provider.phone}', False),
        ('E-mail: {provider.email}', False),
    ])
    fill_cell(t.rows[0].cells[1], [
        ('ЗАКАЗЧИК:', True),
        ('{client.shortName}', True),
        ('', False),
        ('Юридический адрес: {client.legalAddress}', False),
        ('Почтовый адрес: {client.postalAddress}', False),
        ('ИНН: {client.inn}', False),
        ('КПП: {client.kpp}', False),
        ('ОГРН: {client.ogrn}', False),
        ('Расчётный счёт: {client.bankAccount}', False),
        ('{client.bankName}', False),
        ('БИК: {client.bankBik}', False),
        ('Корр. счёт: {client.bankCorrAccount}', False),
        ('Тел.: {client.phone}', False),
        ('E-mail: {client.email}', False),
    ])
    return t


def signatures_table(doc):
    """Двухколоночная таблица подписей."""
    t = doc.add_table(rows=1, cols=2)
    ls = t.rows[0].cells[0]
    rs = t.rows[0].cells[1]
    ls.text = ''
    rs.text = ''

    p = ls.paragraphs[0]
    p.add_run('{provider.signatoryFullName}').bold = True
    ls.add_paragraph()
    ls.add_paragraph('________________ /{provider.signatoryShort}/')
    ls.add_paragraph('м.п.')

    p = rs.paragraphs[0]
    p.add_run('{client.directorRole} {client.shortName}').bold = True
    rs.add_paragraph()
    rs.add_paragraph('________________ /{client.directorName}/')
    rs.add_paragraph('м.п.')
    return t


def price_table_with_loop(doc, headers=None):
    """Таблица прайса с docxtemplater-loop по {#priceItems}...{/priceItems}."""
    if headers is None:
        headers = ['Наименование услуги', 'Площадь, м²', 'Метод обработки',
                   'Кратность', 'Цена, руб. (без НДС)', 'Стоимость, в т.ч. НДС 5%']
    t = doc.add_table(rows=2, cols=len(headers))
    t.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    loop_cells = [
        '{#priceItems}{serviceName}',
        '{area}',
        '{method}',
        '{frequency}',
        '{priceNet}',
        '{priceGross}{/priceItems}',
    ]
    for i, val in enumerate(loop_cells):
        cell = t.rows[1].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
    return t


def objects_table_with_loop(doc):
    """Таблица обработанных объектов: № / Объект / Адрес / Площадь / Услуга."""
    t = doc.add_table(rows=2, cols=5)
    t.style = 'Table Grid'
    headers = ['№', 'Объект', 'Адрес', 'Площадь, м²', 'Услуга']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    loop_cells = [
        '{#objects}{index}',
        '{name}',
        '{address}',
        '{area}',
        '{service}{/objects}',
    ]
    for i, val in enumerate(loop_cells):
        cell = t.rows[1].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
    return t


def object_services_table_with_loop(doc):
    """Sprint 10: таблица услуг объекта для АО/АВР — формат бумажных актов Регины:
    № / Объект / Адрес / Площадь / Услуга. Цикл по {#objectServices} — несколько
    услуг = несколько строк, у каждой своя единица (м²/ед./м³) в колонке «Площадь»."""
    t = doc.add_table(rows=2, cols=5)
    t.style = 'Table Grid'
    headers = ['№', 'Объект', 'Адрес', 'Площадь', 'Услуга']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    loop_cells = [
        '{#objectServices}{index}',
        '{objectName}',
        '{objectAddress}',
        '{areaLabel}',
        '{serviceName}{/objectServices}',
    ]
    for i, val in enumerate(loop_cells):
        cell = t.rows[1].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
    return t


# ═══ 1. ДОГОВОР ═══════════════════════════════════════════════════


def build_contract():
    doc = new_doc()

    # Заголовок
    p = para(doc, '', align='center')
    r = p.add_run('ДОГОВОР НА ОКАЗАНИЕ УСЛУГ № {contract.number}')
    r.bold = True
    r.font.size = Pt(14)
    para(doc)

    p = para(doc)
    p.add_run('Дата заключения: ')
    p.add_run('{contract.date}').bold = True
    p = para(doc)
    p.add_run('Место заключения: ')
    p.add_run('{contract.place}').bold = True
    para(doc)

    # Преамбула
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('{provider.brand}').bold = True
    p.add_run(', в лице ')
    p.add_run('{provider.fullName}').bold = True
    p.add_run(' (далее – Исполнитель), действующий на основании лицензии ')
    p.add_run('{provider.licenseNumber}').bold = True
    p.add_run(' от {provider.licenseDate} года, выданной {provider.licenseAuthority}, '
              'также действующий на основании ОГРНИП: ')
    p.add_run('{provider.ogrnip}').bold = True
    p.add_run(' с одной стороны, и ')
    p.add_run('{client.fullName}').bold = True
    p.add_run(', именуемое в дальнейшем «Заказчик», в лице ')
    p.add_run('{client.directorRole} {client.directorName}').bold = True
    p.add_run(', действующего на основании ')
    p.add_run('{client.actingBasis}').bold = True
    p.add_run(', с другой стороны, вместе именуемые «Стороны», заключили договор о нижеследующем:')

    para(doc)
    heading(doc, '1. Термины, используемые в настоящем договоре')
    para(doc, '1.1. «Дезинсекция» — процедура, в соответствии с которой принимаются санитарные меры по борьбе или '
              'уничтожению насекомых — переносчиков болезней человека.', align='justify')
    para(doc, '1.2. «Дезинфекция» — процедура санитарных мер по борьбе или уничтожению инфекционных агентов '
              'на поверхности тела человека или животного, или же в багаже, грузах, контейнерах, перевозочных '
              'средствах, товарах и почтовых посылках посредством прямого воздействия химических или физических '
              'агентов.', align='justify')
    para(doc, '1.3. «Дератизация» — процедура санитарных мер по борьбе или уничтожению грызунов — переносчиков '
              'болезней человека.', align='justify')

    para(doc)
    heading(doc, '2. Предмет Договора')
    para(doc, '2.1. Заказчик поручает, а Исполнитель принимает обязательство по оказанию услуг (нужное '
              'выделить):', align='justify')
    para(doc, '2.1.1. Дератизация зданий, промышленного оборудования.')
    para(doc, '2.1.2. Дезинсекция.')
    para(doc, '2.1.3. Дезинфекция зданий, промышленного оборудования.')
    para(doc, '2.1.4. Заказчик в свою очередь обязуется указанные услуги оплатить.')
    para(doc, '2.2. Услуги оказываются в объёме, указанном в заявке, направляемой по электронной почте '
              '{provider.email}, либо путём телефонного звонка. Заявка является неотъемлемой частью настоящего '
              'Договора (Приложение № 1).', align='justify')
    para(doc, '2.3. Стоимость услуг определяется согласно Приложению № 2, являющемуся неотъемлемой частью '
              'настоящего Договора.', align='justify')
    para(doc, '2.4. Срок оказания услуг — в течение 2 календарных дней с момента получения заявки и подготовки '
              'помещения.', align='justify')

    para(doc)
    heading(doc, '3. Обязательства Сторон')
    para(doc, '3.1. Заказчик обязуется:', bold=True)
    para(doc, '3.1.1. Обеспечить беспрепятственный доступ специалистов Исполнителя.', align='justify')
    para(doc, '3.1.2. Выполнить подготовку помещений согласно типу услуги.', align='justify')
    para(doc, '3.1.3. Не препятствовать выполнению дезинфекционных мероприятий.', align='justify')
    para(doc, '3.1.4. Назначить ответственное лицо.', align='justify')
    para(doc, '3.1.5. Соблюдать меры безопасности.', align='justify')
    para(doc, '3.1.6. Выполнять рекомендации Исполнителя.', align='justify')
    para(doc, '3.1.7. Оповещать людей на обрабатываемых территориях.', align='justify')
    para(doc, '3.1.8. Оперативно извещать Исполнителя об изменениях.', align='justify')
    para(doc, '3.1.9. По завершении принять результаты и подписать акт приёмки в течение 3 (трёх) рабочих дней.',
         align='justify')
    para(doc, '3.1.10. Оплатить услуги в соответствии со ст. 4 Договора.', align='justify')
    para(doc, '3.2. Заказчик вправе в любое время проверять ход и качество услуг.', align='justify')
    para(doc)
    para(doc, '3.3. Исполнитель обязуется:', bold=True)
    para(doc, '3.3.1. Качественно оказывать услуги в требуемом объёме.', align='justify')
    para(doc, '3.3.2. Проводить мероприятия в соответствии с ФЗ № 52-ФЗ от 30.03.1999, СанПиН 3.3686-21, '
              'СП 1.1.1058-01, МУ-3.4.1127-02.', align='justify')
    para(doc, '3.3.3. Готовить рекомендации по устранению санитарно-технических дефектов.', align='justify')
    para(doc, '3.3.4. Разрабатывать тактику и методику с учётом вида объекта.', align='justify')
    para(doc, '3.3.5. Самостоятельно осуществлять выбор препаратов в соответствии с МУ Минздрава РФ.',
         align='justify')
    para(doc, '3.3.6. Оформить пропуска для оказания услуг в режимных зонах своими силами.', align='justify')
    para(doc, '3.3.7. Передавать акт приёмки услуг уполномоченному лицу Заказчика.', align='justify')
    para(doc, '3.4. В случае отказа Заказчика от приёмки услуг составляется мотивированный двусторонний акт.',
         align='justify')

    para(doc)
    heading(doc, '4. Стоимость услуг и порядок расчётов')
    para(doc, '4.1. Оплата проводится Заказчиком в безналичном порядке по тарифам Приложения № 2. Исполнитель '
              'применяет УСН со ставкой НДС 5%.', align='justify')
    para(doc, '4.2. Оплата услуг должна проводиться в течение 7 (семи) календарных дней со дня получения счёта.',
         align='justify')
    para(doc, '4.3. Датой оплаты признаётся дата зачисления средств на расчётный счёт Исполнителя.',
         align='justify')
    para(doc, '4.4. Возможен авансовый платёж по Договору.', align='justify')
    para(doc, '4.5. Услуги считаются оказанными с момента подписания Сторонами Акта приёмки.', align='justify')
    para(doc, '4.6. При обнаружении отступлений составляется двусторонний протокол с перечнем доработок.',
         align='justify')

    para(doc)
    heading(doc, '5. Ответственность Сторон и разрешение споров')
    para(doc, '5.1. За невыполнение обязательств Стороны несут ответственность в соответствии с '
              'законодательством РФ.', align='justify')
    para(doc, '5.2. За неоказанные услуги Исполнитель уплачивает неустойку 0,1% от стоимости за каждый день.',
         align='justify')
    para(doc, '5.3. За несвоевременную оплату Заказчик уплачивает неустойку 0,1% от стоимости за каждый день.',
         align='justify')
    para(doc, '5.4. Допускается передача документов по электронной почте наравне с письменными сообщениями.',
         align='justify')
    para(doc, '5.5. Допускается использование факсимильного воспроизведения подписи или электронной подписи.',
         align='justify')
    para(doc, '5.6. Исполнитель не несёт ответственности за последствия в случае невыполнения Заказчиком '
              'обязанностей п. 3.1.', align='justify')
    para(doc, '5.7. Если Заказчик не подготовил помещение (п. 3.1.2), Исполнитель имеет право отказаться от '
              'выполнения заявки. Услуга считается оказанной наполовину, оплата 50%.', align='justify')

    para(doc)
    heading(doc, '6. Особые условия')
    para(doc, '6.1. Стороны освобождаются от ответственности за неисполнение обязательств вследствие '
              'непреодолимой силы.', align='justify')
    para(doc, '6.2. Антикоррупционная оговорка: Стороны не предлагали и не принимали неправомерных вознаграждений '
              'в связи с заключением или исполнением настоящего Договора.', align='justify')

    para(doc)
    heading(doc, '7. Порядок урегулирования споров')
    para(doc, '7.1. Спорные вопросы разрешаются переговорами с письменным претензионным порядком.',
         align='justify')
    para(doc, '7.2. Все споры подлежат рассмотрению в Отделении МКАС при ТПП РФ в г. Краснодаре.',
         align='justify')
    para(doc, '7.3. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу.', align='justify')

    para(doc)
    heading(doc, '8. Срок действия Договора')
    p = para(doc)
    p.add_run('8.1. Договор вступает в силу с момента подписания и действует по ')
    p.add_run('{contract.endDate}').bold = True
    p.add_run(', автоматически пролонгируется на следующий календарный год при отсутствии уведомления о '
              'прекращении за один месяц до окончания.')

    para(doc)
    heading(doc, '9. Реквизиты и подписи Сторон')
    para(doc)
    requisites_table(doc)
    para(doc)
    para(doc, 'Подписи сторон', bold=True, align='center')
    para(doc)
    signatures_table(doc)

    # Приложение № 2
    doc.add_page_break()
    para(doc, 'Приложение № 2 к Договору № {contract.number}', align='right', italic=True)
    para(doc, 'от {contract.date}', align='right', italic=True)
    para(doc)
    para(doc, 'СТОИМОСТЬ УСЛУГ', bold=True, align='center', size=14)
    para(doc)
    price_table_with_loop(doc)
    para(doc)
    para(doc)
    signatures_table(doc)

    out = OUT_DIR / 'contract-services.docx'
    doc.save(out)
    return out


# ═══ 2. ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ ═════════════════════════════════


def build_addendum():
    doc = new_doc()

    p = para(doc, '', align='center')
    r = p.add_run('ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ № {addendum.number}')
    r.bold = True
    r.font.size = Pt(13)
    p = para(doc, '', align='center')
    r = p.add_run('к Договору № {contract.number} на оказание услуг от {contract.date}')
    r.bold = True
    r.font.size = Pt(12)
    para(doc)

    p = para(doc)
    p.add_run('Дата заключения: ')
    p.add_run('{addendum.date}').bold = True
    p = para(doc)
    p.add_run('Место заключения: ')
    p.add_run('{addendum.place}').bold = True
    para(doc)

    # Преамбула (краткая)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('{provider.brand}').bold = True
    p.add_run(', в лице ')
    p.add_run('{provider.fullName}').bold = True
    p.add_run(' (далее – Исполнитель), действующий на основании лицензии ')
    p.add_run('{provider.licenseNumber}').bold = True
    p.add_run(' от {provider.licenseDate} года, ОГРНИП: ')
    p.add_run('{provider.ogrnip}').bold = True
    p.add_run(' с одной стороны, и ')
    p.add_run('{client.fullName}').bold = True
    p.add_run(', именуемое «Заказчик», в лице ')
    p.add_run('{client.directorRole} {client.directorName}').bold = True
    p.add_run(', действующего на основании ')
    p.add_run('{client.actingBasis}').bold = True
    p.add_run(', с другой стороны, заключили настоящее Дополнительное соглашение о нижеследующем:')

    para(doc)
    p = para(doc)
    p.add_run('Приложение № 2 к Договору № ').bold = False
    p.add_run('{contract.number}').bold = True
    p.add_run(' на оказание услуг от ').bold = False
    p.add_run('{contract.date}').bold = True
    p.add_run(' изложить в следующей редакции. Утверждённые тарифы (без НДС, со ставкой НДС 5%):')

    para(doc)
    para(doc, 'СТОИМОСТЬ УСЛУГ', bold=True, align='center', size=12)
    para(doc)
    price_table_with_loop(doc)

    para(doc)
    heading(doc, 'Прочие условия', level=2)
    para(doc, '1. Настоящее Дополнительное соглашение вступает в действие в день его подписания Сторонами.',
         align='justify')
    para(doc, '2. Настоящее Дополнительное соглашение является неотъемлемой частью Договора.', align='justify')
    para(doc, '3. Настоящее Дополнительное соглашение составлено в 2 (двух) экземплярах, имеющих одинаковую '
              'юридическую силу, по одному экземпляру для каждой из Сторон.', align='justify')

    para(doc)
    heading(doc, 'Реквизиты и подписи Сторон', level=2)
    para(doc)
    requisites_table(doc)
    para(doc)
    para(doc, 'Подписи сторон', bold=True, align='center')
    para(doc)
    signatures_table(doc)

    out = OUT_DIR / 'agreement-addendum.docx'
    doc.save(out)
    return out


# ═══ 3. АКТ ОБСЛЕДОВАНИЯ ══════════════════════════════════════════


def build_inspection_report():
    doc = new_doc()

    p = para(doc, '', align='center')
    r = p.add_run('АКТ ОБСЛЕДОВАНИЯ № {report.number} от {report.date}')
    r.bold = True
    r.font.size = Pt(14)

    p = para(doc, '', align='center')
    r = p.add_run('на проведение работ по дезинфекции, дезинсекции и дератизации')
    r.italic = True
    r.font.size = Pt(11)
    p = para(doc, '', align='center')
    r = p.add_run('(санитарно-эпидемиологические правила СП 3.3686-21)')
    r.italic = True
    r.font.size = Pt(10)
    para(doc)

    # Реквизиты исполнителя
    para(doc, '{provider.shortName}', bold=True)
    para(doc, 'ИНН: {provider.inn}')
    para(doc, 'ОГРНИП: {provider.ogrnip}')
    para(doc, 'Лицензия: {provider.licenseNumber} от {provider.licenseDate} г.')
    para(doc)

    # Заказчик и объекты
    p = para(doc)
    p.add_run('Заказчик: ')
    p.add_run('{client.shortName}').bold = True
    p = para(doc)
    p.add_run('Договор: № ')
    p.add_run('{contract.number}').bold = True
    p.add_run(' от ')
    p.add_run('{contract.date}').bold = True

    # Контактное лицо на объекте (из карточки объекта) — добавлено по запросу,
    # чтобы в АО был ответственный на месте, как в АВР и заказ-наряде.
    p = para(doc)
    p.add_run('Контактное лицо на объекте: ')
    p.add_run('{contact.fio}').bold = True
    p.add_run(', тел. ')
    p.add_run('{contact.phone}').bold = True

    para(doc)
    object_services_table_with_loop(doc)
    para(doc)

    # Чек-лист состояния
    heading(doc, 'Состояние объекта', level=2)
    p = para(doc)
    p.add_run('Санитарное состояние объекта: ')
    p.add_run('{report.objectStatus}').bold = True
    p.add_run('   (удовлетворительное / неудовлетворительное)')

    p = para(doc)
    p.add_run('Выявленные отклонения: ')
    p.add_run('{report.deviations}').italic = True

    p = para(doc)
    p.add_run('Краткое описание состояния объекта: ')
    p.add_run('{report.description}').italic = True

    p = para(doc)
    p.add_run('Рекомендация: ')
    p.add_run('{report.recommendation}').italic = True

    para(doc)
    p = para(doc)
    p.add_run('Критерии заселённости насекомых/грызунов: ')
    p.add_run('{report.infestationLevel}').bold = True
    p.add_run('   (не заселён / мало заселён / много заселён)')

    p = para(doc)
    p.add_run('Наличие журнала учёта дезинфекционных работ: ')
    p.add_run('{report.hasJournal}').bold = True
    p.add_run('   (нет / да)')

    p = para(doc)
    p.add_run('Состояние журнала учёта: ')
    p.add_run('{report.journalStatus}').bold = True
    p.add_run('   (удовлетворительное / требует замены)')

    para(doc)
    para(doc, 'Подписи сторон', bold=True, align='center')
    para(doc)
    signatures_table(doc)

    out = OUT_DIR / 'inspection-report.docx'
    doc.save(out)
    return out


# ═══ 4. АКТ О ПРИЁМКЕ ВЫПОЛНЕННЫХ РАБОТ ═══════════════════════════


def build_work_completion_report():
    doc = new_doc()

    p = para(doc, '', align='center')
    r = p.add_run('АКТ О ПРИЁМКЕ ВЫПОЛНЕННЫХ РАБОТ № {act.number}')
    r.bold = True
    r.font.size = Pt(14)
    p = para(doc, '', align='center')
    p.add_run('от {act.date}').bold = True
    p = para(doc, '', align='center')
    r = p.add_run('(Лицензия: {provider.licenseNumber} от {provider.licenseDate} г.)')
    r.italic = True
    r.font.size = Pt(10)
    para(doc)

    para(doc, 'Настоящим сообщаем, что была проведена дезинсекция / дезинфекция / дератизация '
              'нижеуказанного объекта в целях обеспечения санитарно-эпидемиологического благополучия '
              'населения по правилам СП 3.3686-21.', align='justify')

    para(doc)
    # Стороны
    t = doc.add_table(rows=1, cols=2)
    fill_cell(t.rows[0].cells[0], [
        ('ЗАКАЗЧИК:', True),
        ('{client.shortName}', True),
        ('ИНН: {client.inn}', False),
        ('Контакт: {act.responsibleName}', False),
        ('Должность: {act.responsibleRole}', False),
        ('Телефон: {act.responsiblePhone}', False),
    ])
    fill_cell(t.rows[0].cells[1], [
        ('ИСПОЛНИТЕЛЬ:', True),
        ('{provider.shortName}', True),
        ('ИНН: {provider.inn}', False),
        ('ОГРНИП: {provider.ogrnip}', False),
        ('Дезинфектор: {act.disinfector}', False),
    ])
    para(doc)

    # Таблица работ (Sprint 10: услуги объекта с единицами)
    object_services_table_with_loop(doc)
    para(doc)

    # Чек проверок
    p = para(doc)
    p.add_run('Качество выполненных работ: ')
    p.add_run('{act.qualityCheck}').bold = True
    p.add_run('   (соответствует / не соответствует)')

    p = para(doc)
    p.add_run('Совпадение фактической площади обработки: ')
    p.add_run('{act.areaCheck}').bold = True
    p.add_run('   (совпадает / не совпадает; фактическая площадь — ')
    p.add_run('{act.actualArea}').bold = True
    p.add_run(' м²)')

    p = para(doc)
    p.add_run('Несоответствие качества работ предъявленным требованиям: ')
    p.add_run('{act.discrepancy}').italic = True

    para(doc)
    heading(doc, 'Общие условия выполнения работ', level=2)
    para(doc, '1. Настоящий АКТ составлен и подписан в двух экземплярах, имеющих равную силу, по одному для '
              'каждой из Сторон.', align='justify')
    para(doc, '2. После подписания АКТа Заказчик подтверждает, что ознакомился и обязуется соблюдать '
              'требования по безопасности, рекомендованные Исполнителем (особенно после работ с химическими '
              'препаратами, клеевыми ловушками и прикормкой для насекомых).', align='justify')
    p = para(doc, '', align='justify')
    p.add_run('3. Настоящий АКТ является неотъемлемой частью Договора № ')
    p.add_run('{contract.number}').bold = True
    p.add_run(' от ')
    p.add_run('{contract.date}').bold = True
    p.add_run(' и содержит перечень выполненных Исполнителем работ.')

    para(doc)
    para(doc, 'Подписи сторон', bold=True, align='center')
    para(doc)
    signatures_table(doc)

    out = OUT_DIR / 'work-completion-report.docx'
    doc.save(out)
    return out


# ═══ 5. КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ ══════════════════════════════════


def build_commercial_offer():
    doc = new_doc()

    # Шапка ДЕЗТЕХЮГ
    p = para(doc, '', align='center')
    r = p.add_run('{provider.brand}')
    r.bold = True
    r.font.size = Pt(20)
    p = para(doc, '', align='center')
    r = p.add_run('Дезинфекционные технологии — ЮГ')
    r.italic = True
    r.font.size = Pt(11)
    p = para(doc, '', align='center')
    p.add_run('{provider.legalAddress}')
    p = para(doc, '', align='center')
    p.add_run('Тел.: {provider.phone}   ·   E-mail: {provider.email}')
    para(doc)

    # Заголовок
    p = para(doc, '', align='center')
    r = p.add_run('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ № {offer.number}')
    r.bold = True
    r.font.size = Pt(14)
    p = para(doc, '', align='center')
    p.add_run('от {offer.date}').bold = True
    para(doc)

    # Адресат
    p = para(doc)
    p.add_run('Кому: ').bold = True
    p.add_run('{client.shortName}').bold = True
    p = para(doc)
    p.add_run('Внимание: ')
    p.add_run('{client.directorRole} {client.directorName}').bold = True
    para(doc)

    # Вступление
    para(doc, 'Уважаемый партнёр!', bold=True)
    para(doc, '{offer.intro}', align='justify')
    para(doc, 'ИП Белавина О.В. — лицензированный исполнитель санитарных услуг по Краснодарскому краю '
              '(лицензия {provider.licenseNumber} от {provider.licenseDate}). Применяем '
              'современные сертифицированные препараты и работаем с объектами любой сложности — от '
              'квартир и офисов до промышленных складов и сетей общественного питания.', align='justify')

    para(doc)
    para(doc, 'Предлагаем следующий перечень услуг:', bold=True)
    para(doc)
    price_table_with_loop(doc)

    para(doc)
    p = para(doc, '', align='right')
    p.add_run('Итого без НДС: ').bold = True
    p.add_run('{offer.totalNet}').bold = True
    p.add_run(' руб.')
    p = para(doc, '', align='right')
    p.add_run('Стоимость с НДС 5%: ').bold = True
    p.add_run('{offer.totalGross}').bold = True
    p.add_run(' руб.')
    p = para(doc, '', align='right')
    p.add_run('Сумма прописью: ')
    p.add_run('{offer.totalInWords}').italic = True

    para(doc)
    p = para(doc)
    p.add_run('Срок действия предложения: до ')
    p.add_run('{offer.validUntil}').bold = True
    p.add_run('. По истечении срока тарифы могут быть пересмотрены.')

    para(doc, 'Готовы выехать на бесплатный замер объекта в удобное для вас время. По вопросам '
              'согласования и подписания договора обращайтесь по контактам выше.', align='justify')

    para(doc)
    para(doc, 'С уважением,', italic=True)
    para(doc, '{provider.signatoryFullName}', bold=True)
    para(doc, '________________ /{provider.signatoryShort}/')
    para(doc, 'м.п.')

    out = OUT_DIR / 'commercial-offer.docx'
    doc.save(out)
    return out


# ═══ 6. СЧЁТ НА ОПЛАТУ ════════════════════════════════════════════


def build_invoice():
    doc = new_doc()

    # Шапка с банковскими реквизитами (как в стандартном счёте)
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'

    # Строка 1: Банк получателя
    fill_cell(t.rows[0].cells[0], [('Банк получателя', False)])
    fill_cell(t.rows[0].cells[1], [('{provider.bankName}', True)])
    # Строка 2: БИК + Корр счёт
    fill_cell(t.rows[1].cells[0], [('БИК', False), ('Сч. №', False)])
    fill_cell(t.rows[1].cells[1], [('{provider.bankBik}', True), ('{provider.bankCorrAccount}', True)])
    # Строка 3: Получатель
    fill_cell(t.rows[2].cells[0], [
        ('ИНН: {provider.inn}', False),
        ('Получатель', False),
    ])
    fill_cell(t.rows[2].cells[1], [
        ('ОГРНИП: {provider.ogrnip}', False),
        ('{provider.fullName}', True),
    ])
    # Строка 4: Расчётный счёт
    fill_cell(t.rows[3].cells[0], [('Сч. № получателя', False)])
    fill_cell(t.rows[3].cells[1], [('{provider.bankAccount}', True)])

    para(doc)

    # Заголовок счёта
    p = para(doc, '', align='center')
    r = p.add_run('СЧЁТ НА ОПЛАТУ № {invoice.number} от {invoice.date}')
    r.bold = True
    r.font.size = Pt(14)
    para(doc)

    # Поставщик / Покупатель
    p = para(doc)
    p.add_run('Поставщик (Исполнитель): ').bold = True
    p.add_run('{provider.shortName}, ИНН {provider.inn}, ОГРНИП {provider.ogrnip}, '
              '{provider.legalAddress}, тел. {provider.phone}')

    p = para(doc)
    p.add_run('Покупатель (Заказчик): ').bold = True
    p.add_run('{client.shortName}, ИНН {client.inn}, КПП {client.kpp}, '
              '{client.legalAddress}, тел. {client.phone}')

    p = para(doc)
    p.add_run('Основание: ').bold = True
    p.add_run('{invoice.basis}').italic = True
    p.add_run('   (например: Договор № {contract.number} от {contract.date})')

    para(doc)

    # Таблица позиций
    items_table = doc.add_table(rows=2, cols=6)
    items_table.style = 'Table Grid'
    headers = ['№', 'Наименование товара / услуги', 'Кол-во', 'Ед.', 'Цена', 'Сумма']
    for i, h in enumerate(headers):
        cell = items_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    loop_cells = [
        '{#priceItems}{index}',
        '{serviceName}',
        '{quantity}',
        '{unit}',
        '{priceNet}',
        '{amount}{/priceItems}',
    ]
    for i, val in enumerate(loop_cells):
        cell = items_table.rows[1].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)

    para(doc)
    p = para(doc, '', align='right')
    p.add_run('Итого: ').bold = True
    p.add_run('{invoice.totalNet}').bold = True
    p.add_run(' руб.')
    p = para(doc, '', align='right')
    p.add_run('В т.ч. НДС 5%: ')
    p.add_run('{invoice.vatAmount}').bold = True
    p.add_run(' руб.')
    p = para(doc, '', align='right')
    p.add_run('Всего к оплате: ').bold = True
    p.add_run('{invoice.totalGross}').bold = True
    p.add_run(' руб.')

    para(doc)
    p = para(doc)
    p.add_run('Сумма прописью: ').bold = True
    p.add_run('{invoice.totalInWords}').italic = True

    para(doc)
    p = para(doc)
    p.add_run('Срок оплаты: ')
    p.add_run('до {invoice.dueDate}').bold = True
    p.add_run('. Оплата производится в безналичном порядке на расчётный счёт Поставщика.')

    para(doc)
    para(doc, 'Внимание! Оплата данного счёта означает согласие с условиями поставки услуг.',
         italic=True, size=10)
    para(doc, 'Уведомление об оплате обязательно направить по адресу: {provider.email}', italic=True, size=10)

    para(doc)
    para(doc)
    p = para(doc)
    p.add_run('Поставщик / Исполнитель: ').bold = True
    p.add_run('{provider.signatoryFullName}').bold = True
    para(doc, '________________ /{provider.signatoryShort}/    м.п.')

    out = OUT_DIR / 'invoice.docx'
    doc.save(out)
    return out


# ─── Main ────────────────────────────────────────────────────────


BUILDERS = {
    'contract': build_contract,
    'addendum': build_addendum,
    'inspection': build_inspection_report,
    'work': build_work_completion_report,
    'offer': build_commercial_offer,
    'invoice': build_invoice,
}


def main():
    targets = sys.argv[1:] or list(BUILDERS.keys())
    for t in targets:
        if t not in BUILDERS:
            print(f'  ! unknown target: {t} (available: {list(BUILDERS.keys())})')
            continue
        out = BUILDERS[t]()
        size = out.stat().st_size
        print(f'  OK  {t:12s} -> {out} ({size:,} bytes)')


if __name__ == '__main__':
    main()
