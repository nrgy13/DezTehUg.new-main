#!/usr/bin/env python3
"""
Сборка docxtemplater-шаблонов «точь-в-точь» из РЕАЛЬНЫХ доков ИП Белавина.

Идея: берём реальный .docx компании и хирургически заменяем конкретику на
{плейсхолдеры} build-data.ts, СОХРАНЯЯ форматирование, логотип, QR, подписи.
Это НЕ python-docx-генерация с нуля (её не повторить 1:1), а правка живого файла.

Источники (gitignored, лежат у Сани локально):
  DTUdocs/ — акты/ДС ИП Белавина
  DTU_docs_clients/Договора на заключение — копия/ — клиентские договоры
Выходы: templates/*.docx (активный seed читает их, см. lib/templates/get-template.ts).

ЗАПУСК:  python tools/build-templates-from-real.py ao   (ao|avr|dogovor|ds|all)
Зависимость: pip install python-docx. Валидация рендера: tools/render-test-template.js.

═══ КРИТИЧЕСКАЯ ГРАБЛЯ ═══
python-docx `run.text = ""` УДАЛЯЕТ ВСЁ содержимое run'а, включая <w:drawing>
(логотип/QR живут в run'ах!). Поэтому set_merge() ПРОПУСКАЕТ run'ы с картинками —
чистит только текстовые. Проверка: <a:blip> в выводе должно совпадать с исходником.

═══ ПЛЕЙСХОЛДЕРЫ (из build-data.ts) ═══
provider.* (ИП Белавина, фикс — можно статикой); client.shortName/fullName/
directorName/directorShort(инициалы)/directorRole/actingBasis/legalAddress/
postalAddress/inn/kpp/ogrn/phone/email/bankName/bankAccount/bankBik/bankCorrAccount;
contract.number/date/place/endDate; addendum.number/date/place;
act.number/date/disinfector/responsibleName/responsibleRole/responsiblePhone
(qualityCheck/areaCheck — НЕ используем в АВР: выбор руками, решение Сани);
report.date (АО: только дата + таблица; чекбоксы руками);
contact.fio/phone; object.name/address/area; totalNet/totalGross/vatAmount;
Циклы: {#objectServices}{index}{objectName}{objectAddress}{areaLabel}{serviceName}{/objectServices}
       {#priceItems}{objectName}{serviceName}{area}{priceGross}{vatLine}{frequency}{objectAddress}{/priceItems}
       {#priceItemsByObject}{objectName}{objectAddress}{#items}...{/items}{/priceItemsByObject}  (ДС, мульти-объект)
"""
import copy
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

ROOT = __import__('os').path.abspath(__import__('os').path.join(__import__('os').path.dirname(__file__), '..'))


def has_drawing(run):
    return run._r.find(qn('w:drawing')) is not None or run._r.find(qn('w:pict')) is not None


def full_text(p):
    return "".join(r.text for r in p.runs if not has_drawing(r))


def set_merge(p, new_text):
    """Записать new_text в первый ТЕКСТОВЫЙ run, остальные текстовые очистить.
    Run'ы с картинками (drawing) НЕ трогаем — иначе теряем логотип/QR."""
    text_runs = [r for r in p.runs if not has_drawing(r)]
    if not text_runs:
        p.add_run(new_text)
        return
    text_runs[0].text = new_text
    for r in text_runs[1:]:
        r.text = ""


def replace_run(p, needle, new, whole=False):
    """Точечная замена в ПЕРВОМ текстовом run, содержащем needle.
    whole=True — заменить весь run на new (для подчёркиваний/дат с переменным числом
    пробелов); whole=False — заменить только подстроку needle→new.
    Картиночные run'ы (логотип/QR) НЕ трогаем. Возвращает True, если нашёл."""
    for r in p.runs:
        if has_drawing(r):
            continue
        if needle in r.text:
            r.text = new if whole else r.text.replace(needle, new)
            return True
    return False


def append_run(p, text):
    """Дописать run с text в конец параграфа, скопировав формат (rPr) первого
    значащего текстового run — чтобы шрифт/размер совпадали с меткой строки."""
    new_r = p.add_run(text)
    for r in p.runs:
        if r is new_r or has_drawing(r):
            continue
        if r.text.strip():
            rpr = r._r.find(qn('w:rPr'))
            if rpr is not None:
                new_r._r.insert(0, copy.deepcopy(rpr))
            break
    return new_r


def set_paragraph_parts(p, parts):
    """Пересоздать параграф из частей [(text, bold)], сохранив базовый формат
    (size/font) первого текстового run, но переопределив жирность по каждой части.
    Нужно когда set_merge схлопнул бы смешанное форматирование (напр. преамбула,
    где жирное только название компании)."""
    base_rpr = None
    for r in p.runs:
        if not has_drawing(r) and r.text.strip():
            el = r._r.find(qn('w:rPr'))
            if el is not None:
                base_rpr = copy.deepcopy(el)
            break
    for r in list(p.runs):
        if not has_drawing(r):
            r._r.getparent().remove(r._r)
    for text, bold in parts:
        run = p.add_run(text)
        if base_rpr is not None:
            run._r.insert(0, copy.deepcopy(base_rpr))
        run.bold = bold


def table_to_object_cycle(t, header_rows=1):
    """Превратить таблицу №|Объект|Адрес|Площадь|Услуга в один цикл-ряд {#objectServices}.
    Оставляет header_rows строк-заголовков, первую строку данных делает циклом,
    остальные строки данных удаляет."""
    data_row = t.rows[header_rows]
    vals = ["{#objectServices}{index}", "{objectName}", "{objectAddress}",
            "{areaLabel}", "{serviceName}{/objectServices}"]
    for c, v in zip(data_row.cells, vals):
        set_merge(c.paragraphs[0], v)
        for extra in c.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
    for tr in [r._tr for r in t.rows[header_rows + 1:]]:
        tr.getparent().remove(tr)


def table_to_price_cycle(t, header_rows=1):
    """Прайс-таблицу Наименование|Услуга|Площадь|Цена|Кратность|Адрес → цикл {#priceItems}.
    Оставляет header, первую строку данных делает циклом, остальные удаляет."""
    data_row = t.rows[header_rows]
    vals = ["{#priceItems}{objectName}", "{serviceName}", "{area}",
            "{priceGross}", "{frequency}", "{objectAddress}{/priceItems}"]
    for c, v in zip(data_row.cells, vals):
        set_merge(c.paragraphs[0], v)
        for extra in c.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
    for tr in [r._tr for r in t.rows[header_rows + 1:]]:
        tr.getparent().remove(tr)


def table_to_items_cycle(t, header_rows=1):
    """Прайс-таблица ДС → ВЛОЖЕННЫЙ цикл {#items} (внутри {#priceItemsByObject}).
    Колонки: Наименование|Площадь|без НДС|НДС 5%|с НДС|Периодичность."""
    data_row = t.rows[header_rows]
    vals = ["{#items}{serviceName}", "{area}", "{priceNet}",
            "{vatLine}", "{priceGross}", "{frequency}{/items}"]
    for c, v in zip(data_row.cells, vals):
        set_merge(c.paragraphs[0], v)
        for extra in c.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
    for tr in [r._tr for r in t.rows[header_rows + 1:]]:
        tr.getparent().remove(tr)


# Инициалы в строке подписи: «/А.А.Дёмин/», «/А.Е.Мороз/» → «/{client.directorShort}/».
SIG_INITIALS = re.compile(r'/\s*[А-ЯЁ]\.[А-ЯЁ]\.[А-ЯЁ][а-яё]+\s*/')


def fill_signature_cell(cell):
    """Блок подписи Заказчика (правая ячейка): строка с инициалами «/Х.Х.Фамилия/» →
    «/{client.directorShort}/» (подчёркивания сохраняем); строка с ролью+названием →
    «{client.directorRole} {client.shortName}». Универсально для договора и ДС."""
    for p in cell.paragraphs:
        t = full_text(p)
        if not t.strip():
            continue
        if SIG_INITIALS.search(t):
            set_merge(p, SIG_INITIALS.sub('/{client.directorShort}/', t))
        elif "директор" in t.lower():
            set_merge(p, "{client.directorRole} {client.shortName}")


def set_table_cell_margins(table, top=60, bottom=60, left=120, right=120):
    """Задать внутренние отступы ячеек таблицы (в twips) — чтобы текст не лип к границам.
    Действует на всю таблицу через tblPr/tblCellMar (сохраняется при повторе цикла)."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)


def space_before_paragraph(p, pts=18):
    """Добавить вертикальный отступ перед параграфом (для «Подписи сторон» и т.п.)."""
    p.paragraph_format.space_before = Pt(pts)


def finalize_signatures(t, empty=2):
    """Привести ОБЕ ячейки подписи к одинаковому жёсткому виду:
    [строка имени/роли][empty пустых][строка подписи /Фамилия/].
    Так и верхняя строка, и подпись стоят на одном уровне в обеих колонках,
    независимо от исходной (разной в TABLE 1/3/6) структуры ячеек."""
    for c in t.rows[0].cells:
        title = next((p for p in c.paragraphs
                      if full_text(p).strip() and '____' not in full_text(p)), None)
        sign = next((p for p in c.paragraphs if '____' in full_text(p)), None)
        if title is None or sign is None:
            continue
        # снять переносы и в заголовке (TABLE6: «\nИП Белавина» — ведущий перенос сдвигал
        # имя вниз), и в строке подписи (set_merge мог оставить <w:br>/буквальный \n)
        for p in (title, sign):
            for r in p.runs:
                for br in r._r.findall(qn('w:br')):
                    r._r.remove(br)
                if '\n' in r.text:
                    r.text = r.text.replace('\n', '')
        # оставить только заголовок и подпись, убрав весь остальной (пустой) мусор
        keep = {title._p, sign._p}
        for p in list(c.paragraphs):
            if p._p not in keep:
                p._p.getparent().remove(p._p)
        # вставить ровно `empty` пустых параграфов между заголовком и подписью
        for _ in range(empty):
            sign.insert_paragraph_before('')


# Реквизиты Заказчика (правая ячейка TABLE 0): якорь в исходной строке → новый текст параграфа.
# Порядок важен — «Расчетный счет»/«Корр. счет» различаем полными словами.
REQ_RULES = [
    ("ЗАКАЗЧИК", "ЗАКАЗЧИК:"),
    ("Юридический адрес", "Юридический адрес: {client.legalAddress}"),
    ("Фактический адрес", "Фактический адрес: {client.postalAddress}"),
    ("Расчетный счет", "Расчетный счет: {client.bankAccount}"),
    ("Корр. счет", "Корр. счет: {client.bankCorrAccount}"),
    ("БИК", "БИК: {client.bankBik}"),
    ("ОГРН", "ОГРН {client.ogrn}"),
    ("ИНН", "ИНН {client.inn} КПП {client.kpp}"),
    ("Сбербанк", "{client.bankName}"),
    ("Тел.", "Тел. {client.phone}"),
    ("Е-mail", "Е-mail: {client.email}"),
    ("АРУМ СЕРВИС", "{client.shortName}"),
]

# Реквизиты Заказчика в ДС (формат ИП Белавина чуть другой: «Почтовый адрес», КПП и ОГРН
# отдельными строками, двоеточия после меток). Якорь в исходной строке → новый текст.
DS_REQ_RULES = [
    ("ЗАКАЗЧИК", "ЗАКАЗЧИК:"),
    ("Юридический адрес", "Юридический адрес: {client.legalAddress}"),
    ("Почтовый адрес", "Почтовый адрес: {client.postalAddress}"),
    ("Расчетный счет", "Расчетный счет: {client.bankAccount}"),
    ("Корр. счет", "Корр. счет: {client.bankCorrAccount}"),
    ("БИК", "БИК: {client.bankBik}"),
    ("ОГРН", "ОГРН: {client.ogrn}"),
    ("КПП", "КПП: {client.kpp}"),
    ("ИНН", "ИНН: {client.inn}"),
    ("СБЕРБАНК", "{client.bankName}"),
    ("Тел", "Тел.: {client.phone}"),
    ("Е-mail", "Е-mail: {client.email}"),
    ("Аппетит", "{client.shortName}"),
]


def fill_requisites_cell(cell, rules):
    """Заменить реквизиты ПОСТРОЧНО, сохраняя разрывы строк <w:br/>.
    ⚠️ Реквизиты бывают двух форматов: отдельный параграф на строку (договор) ИЛИ
    несколько строк в одном параграфе через <w:br/> (ДС). Поэтому режем по \\n (текст
    run'а включает <w:br/> как \\n), применяем первое подходящее правило к КАЖДОЙ строке
    и пересобираем параграф с теми же разрывами — иначе set_merge на весь параграф
    затирает соседние строки."""
    for p in cell.paragraphs:
        full = "".join(r.text for r in p.runs)
        if not full.strip():
            continue
        new_lines = []
        for line in full.split("\n"):
            repl = line
            if line.strip():
                for marker, new in rules:
                    if marker in line:
                        repl = new
                        break
            new_lines.append(repl)
        # сохранить формат (rPr) первого run, затем пересоздать run'ы + разрывы
        rpr = None
        for r in p.runs:
            el = r._r.find(qn('w:rPr'))
            if el is not None:
                rpr = copy.deepcopy(el)
                break
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        for i, line in enumerate(new_lines):
            run = p.add_run(line)
            if rpr is not None:
                run._r.insert(0, copy.deepcopy(rpr))
            if i < len(new_lines) - 1:
                run.add_break()


# ─────────────────────────────── АО (act_inspection) ✅ ГОТОВ ───────────────────────────────
def build_ao():
    src = f"{ROOT}/DTUdocs/Акт обледования шаблон ИП Белавина.docx"
    out = f"{ROOT}/templates/inspection-report.docx"
    doc = Document(src)
    # 1) Заголовок «АКТ-ОБСЛЕДОВАНИЯ от ___» → + {report.date}
    for p in doc.paragraphs:
        t = full_text(p)
        if "АКТ" in t and "ОБСЛЕДОВАНИ" in t:
            nt = re.sub(r"от\s*$", "от {report.date}", t.rstrip())
            if "{report.date}" not in nt:
                nt = nt + " {report.date}"
            set_merge(p, nt)
            break
    # 2) Строка контакта перед «Санитарное состояние» (после таблицы) — пожелание Сани
    for p in doc.paragraphs:
        if "Санитарное состояние" in full_text(p):
            newp = p.insert_paragraph_before("Контактное лицо на объекте: {contact.fio}, тел. {contact.phone}")
            for r in newp.runs:
                r.font.name = "Times New Roman"; r.font.size = Pt(12)
            break
    # 3) Таблица объектов (header + 4 пустых) → один цикл-ряд
    t = doc.tables[0]
    vals = ["{#objectServices}{index}", "{objectName}", "{objectAddress}", "{areaLabel}", "{serviceName}{/objectServices}"]
    for c, v in list(zip(t.rows[1].cells, vals)):
        set_merge(c.paragraphs[0], v)
        for extra in c.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
    for tr in [r._tr for r in t.rows[2:]]:
        tr.getparent().remove(tr)
    doc.save(out)
    print("OK ao ->", out)


# ─────────────────────────────── АВР (act_work) ✅ ГОТОВ ───────────────────────────────
def build_avr():
    """Источник: DTUdocs/Акт по проведению работ ИП Белавина О.В. Столовая СОК Анапа.docx
       → templates/work-completion-report.docx
    Структура (inspect-docx): blip=2 (лого+QR в заголовке P0!), 38 параграфов, 1 таблица 3x5.
    Решения Сани: соответствует/совпадает — статикой (выбор руками); уполномоченное лицо — авто.
    Форматы (build-data): act.date/contract.date = «7 апреля 2026 г.» (С « г.»); act.number = офиц. номер."""
    src = f"{ROOT}/DTUdocs/Акт по проведению работ ИП Белавина О.В. Столовая СОК Анапа.docx"
    out = f"{ROOT}/templates/work-completion-report.docx"
    doc = Document(src)
    P = doc.paragraphs

    def find(substr):
        for p in P:
            if substr in full_text(p):
                return p
        return None

    # P0 — заголовок (содержит лого+QR!): № под номер + дата. Только точечная замена run'ов.
    p0 = find("АКТ О ПРИЕМКЕ")
    replace_run(p0, "_____", "№{act.number} ", whole=True)  # «№ _____ _» → «№{act.number} » (пробел перед «от»)
    replace_run(p0, "30.03.2026", "{act.date}", whole=True)  # «30.03.2026 г.» → «{act.date}» (уже с «г.»)

    # Заказчик (обычный) + ИНН (метка жирная, значение обычное — как в оригинале)
    set_merge(find("«Аппетит»"), "{client.shortName}")
    set_paragraph_parts(find("ИНН: 2320155529"), [("ИНН: ", True), ("{client.inn}", False)])

    # Контакт заказчика (есть переносы \n → точечно)
    p_ct = find("Александр Витальевич Сафронов")
    replace_run(p_ct, "Александр Витальевич Сафронов", "{act.responsibleName}", whole=True)
    replace_run(p_ct, "8-904-492-54-63", "{act.responsiblePhone}", whole=True)

    # Дезинфектор (блок ИСПОЛНИТЕЛЯ, есть переносы → точечно; ИНН/ОГРНИП Белавиной — статика)
    p_di = find("Дезинфектор:")
    replace_run(p_di, "Нечепоренко", "{act.disinfector}", whole=True)
    replace_run(p_di, "Д.И.", "", whole=True)

    # Ссылка на договор — собираем заново (не задвоить год/«г.»); № и дата жирные (как в оригинале)
    set_paragraph_parts(find("является неотъемлемой частью Договора"), [
        ("2. Настоящий АКТ является неотъемлемой частью Договора ", False),
        ("№", True),
        ("{contract.number} ", False),
        ("от {contract.date}", True),
        (", оформлен в установленном порядке, содержит перечень выполняемых работ, "
         "предоставленных «Исполнителем»", False),
    ])

    # Уполномоченное лицо (АВТО, решение Сани) — «метка: значение» вместо прочерка-табуляции
    # (иначе значение прилипало в конец линии и длинный телефон переносился на 2 строки)
    set_merge(find("Фамилия Имя Отчество"), "Фамилия Имя Отчество: {act.responsibleName}")
    set_merge(find("Должность"), "Должность: {act.responsibleRole}")
    set_merge(find("Телефон"), "Телефон: {act.responsiblePhone}")

    # Таблица объектов/услуг → цикл {#objectServices} (как в АО)
    table_to_object_cycle(doc.tables[0], header_rows=1)

    doc.save(out)
    print("OK avr ->", out)


# ─────────────────────────────── Договор (contract) ✅ ГОТОВ ───────────────────────────────
def build_dogovor():
    """Источник: DTU_docs_clients/Договора на заключение — копия/Договор ООО АРУМ СЕРВИС от 01.01.2026.docx
       → templates/contract-services.docx
    Структура (inspect-docx): blip=0, 172 параграфа, 7 таблиц.
      TABLE0 = реквизиты сторон (правая=Заказчик), TABLE1/3/6 = подписи (правая=Заказчик),
      TABLE2 = бланк заявки (статика), TABLE4 (46x6) = прайс клиента → {#priceItems},
      TABLE5 = прочие услуги прейскурант (статика).
    СТАТИКА: преамбула Исполнителя (P7), юр.текст 1-8, заявка, прочие услуги.
    ⚠️ contract.endDate может быть пуст → «действует по , если» (fallback в build-data добавлен)."""
    src = (f"{ROOT}/DTU_docs_clients/Договора на заключение — копия/"
           "Договор ООО АРУМ СЕРВИС от 01.01.2026.docx")
    out = f"{ROOT}/templates/contract-services.docx"
    doc = Document(src)
    P = doc.paragraphs

    def find(substr):
        for p in P:
            if substr in full_text(p):
                return p
        return None

    # Преамбула: номер / дата / место
    set_merge(find("ДОГОВОР НА ОКАЗАНИЕ УСЛУГ"), "ДОГОВОР НА ОКАЗАНИЕ УСЛУГ №{contract.number}")
    set_merge(find("Дата заключения"), "Дата заключения: {contract.date}")
    set_merge(find("Место заключения"), "Место заключения: {contract.place}")
    # Вводный абзац Заказчика (статика Исполнителя выше — НЕ трогаем).
    # Название компании — жирное (как в оригинале), остальное обычное.
    set_paragraph_parts(find("заключили договор о нижеследующем"), [
        ("и ", False),
        ("{client.fullName}", True),
        (", в лице {client.directorRole} {client.directorName}, действующего на основании "
         "{client.actingBasis}, именуемое в дальнейшем «Заказчик» с другой стороны, "
         "вместе именуемые «Стороны», заключили договор о нижеследующем:", False),
    ])
    # Срок действия (8.1)
    set_merge(
        find("Договор вступает в силу"),
        "8.1.\tДоговор вступает в силу с момента его подписания и действует по {contract.endDate}, "
        "если ни одна из Сторон не заявят о прекращении его действия за один месяц до окончания "
        "календарного года, Договор считается пролонгированным на прежних условиях и на очередной "
        "календарный год. Количество продлений действия Договора не ограничивается. Исполнитель имеет "
        "право не чаще чем один раз в квартал пересматривать условия положения Приложения № 2 "
        "настоящего Договора, о чем он обязан письменно уведомить Заказчика.",
    )
    # Приложения №1/№2 — номер договора
    set_merge(find("Приложение № 1 к Договору"), "Приложение № 1 к Договору №{contract.number}")
    set_merge(find("Приложение № 2 к Договору"), "Приложение № 2 к Договору №{contract.number}")
    # Даты приложений («от 01 января 2026 г», оба = дата договора)
    for p in P:
        if full_text(p).strip().startswith("от 01 января"):
            set_merge(p, "от {contract.date}")
    # Заголовок прайса — убрать конкретный объект (артефакт договора АРУМ) + лишние
    # пустые строки от него; оставить просто «СТОИМОСТЬ УСЛУГ»
    set_merge(find("СТОИМОСТЬ УСЛУГ"), "СТОИМОСТЬ УСЛУГ")
    # Дата начала работ
    set_merge(
        find("Дата начала работ"),
        "Дата начала работ по дезинсекции, дезинфекции и дератизации объекта - {contract.date}",
    )

    # Реквизиты Заказчика (TABLE0, правая ячейка)
    fill_requisites_cell(doc.tables[0].rows[0].cells[1], REQ_RULES)
    # Подписи Заказчика (TABLE 1/3/6, правая ячейка) + выравнивание подписей по низу
    for ti in (1, 3, 6):
        fill_signature_cell(doc.tables[ti].rows[0].cells[1])
        finalize_signatures(doc.tables[ti])
    # Прайс клиента (TABLE4) → цикл {#priceItems}
    table_to_price_cycle(doc.tables[4])
    # Отступ перед заголовками «Подписи сторон» (а то прижато к таблице выше)
    for p in doc.paragraphs:
        if full_text(p).strip() == "Подписи сторон":
            space_before_paragraph(p, 18)

    doc.save(out)
    print("OK dogovor ->", out)


# ─────────────────────────────── ДС (addendum) ✅ ГОТОВ ───────────────────────────────
def build_ds():
    """Источник: DTUdocs/ДС№2 ООО Аппетит.docx → templates/agreement-addendum.docx
    Структура (inspect-docx): blip=0, 24 параграфа, 4 таблицы.
      TABLE0/1 (3x6) = прайс по объекту 1/2 (мульти-объект), TABLE2 (2x2) = реквизиты,
      TABLE3 (1x2) = подписи. P12/P13 = заголовки-адреса объектов.
    МУЛЬТИ-ОБЪЕКТ (решение Сани): внешний цикл {#priceItemsByObject} оборачивает
      [заголовок-абзац объекта + прайс-таблицу]; внутри таблицы — {#items}.
      P12 открывает цикл, P13 закрывает, TABLE1 (дубль) удаляется — docxtemplater
      повторит «заголовок+таблицу» для каждого объекта."""
    src = f"{ROOT}/DTUdocs/ДС№2 ООО Аппетит.docx"
    out = f"{ROOT}/templates/agreement-addendum.docx"
    doc = Document(src)
    P = doc.paragraphs

    def find(substr):
        for p in P:
            if substr in full_text(p):
                return p
        return None

    # Шапка ДС: номер ДС + номер договора + даты + место
    set_merge(find("Дополнительное соглашение"),
              "Дополнительное соглашение №{addendum.number} к ДОГОВОРУ №{contract.number}")
    set_merge(find("на оказание услуг от"), "на оказание услуг от {contract.date}")
    set_merge(find("Дата заключения"), "Дата заключения: {addendum.date}")
    set_merge(find("Место заключения"), "Место заключения: {addendum.place}")
    # Преамбула Заказчика (точь-в-точь ДС — без «в лице директора», он в подписи).
    # Название компании и слово «Заказчик» — жирные (как в оригинале).
    set_paragraph_parts(find("заключили договор о нижеследующем"), [
        ("и ", False),
        ("{client.fullName}", True),
        (", действующего на основании {client.actingBasis}, именуемое в дальнейшем ", False),
        ("«Заказчик»", True),
        (" с другой стороны, вместе именуемые «Стороны», заключили договор о нижеследующем:", False),
    ])
    # Отсылка к Приложению №2 договора (номер договора + дата договора)
    set_merge(find("дополнить следующими данными"),
              "Приложение №2 к Договору №{contract.number} на оказание услуг по проведению "
              "дезинсекционных мероприятий от {contract.date} дополнить следующими данными. "
              "Исполнитель применяет УСН со ставкой НДС-5%:")

    # Ссылки на таблицы ДО структурных правок (индексы сдвинутся после удаления)
    t_price0, t_price1, t_req, t_sign = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]

    # МУЛЬТИ-ОБЪЕКТ: открывающий/закрывающий теги — в ОТДЕЛЬНЫХ параграфах, а заголовок
    # объекта внутри цикла одним параграфом. Иначе docxtemplater тащит формат P12 на 1-й
    # объект, P13 — на 2-й (заголовки получались разного вида). Пустой параграф открытия
    # даёт ещё и отступ между блоками объектов.
    p_obj = find("столовая СОК Анапа Нептун")  # P12 — формат заголовка (жирный, центр)
    p_obj.insert_paragraph_before("{#priceItemsByObject}")
    set_merge(p_obj, "{objectName} ({objectAddress})")
    set_merge(find("Курортная Деревня"), "{/priceItemsByObject}")
    # TABLE0 — вложенный цикл строк {#items}; TABLE1 (дубль 2-го объекта) удаляем
    table_to_items_cycle(t_price0)
    set_table_cell_margins(t_price0, top=60, bottom=60, left=120, right=120)  # воздух в ячейках
    t_price1._tbl.getparent().remove(t_price1._tbl)

    # Реквизиты Заказчика (правая ячейка) + подпись Заказчика (выровнять по низу)
    fill_requisites_cell(t_req.rows[0].cells[1], DS_REQ_RULES)
    fill_signature_cell(t_sign.rows[0].cells[1])
    finalize_signatures(t_sign)
    # Отступ перед «Подписи сторон»
    for p in doc.paragraphs:
        if full_text(p).strip() == "Подписи сторон":
            space_before_paragraph(p, 18)

    doc.save(out)
    print("OK ds ->", out)


BUILDERS = {"ao": build_ao, "avr": build_avr, "dogovor": build_dogovor, "ds": build_ds}

if __name__ == "__main__":
    arg = sys.argv[1] if len(sys.argv) > 1 else "ao"
    targets = list(BUILDERS) if arg == "all" else [arg]
    for t in targets:
        if t not in BUILDERS:
            print("Неизвестная цель:", t, "— доступно:", list(BUILDERS)); sys.exit(1)
        BUILDERS[t]()
